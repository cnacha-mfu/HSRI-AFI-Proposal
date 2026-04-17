# -*- coding: utf-8 -*-
"""
Update Category 1:
  - PI back to 20% FTE (162,000)
  - Co-PI back to 15% FTE (94,500)
  - AI/ML Dev 1 & 2 label corrected to Master's (22,000 unchanged)
  - Clinical RA: Bachelor's 18,000/mo, expanded to 2 persons
Recalculates Cat 5 fee and grand total.
"""
import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)
t   = doc.tables[6]

def shade(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, italic=False, size=10,
             align='left', fill=None, color_hex=None):
    if fill:
        shade(cell, fill)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT  if align == 'right'  else
                      WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else
                      WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(r, g, b)

def update_item_row(row, name, detail, amount, fill):
    set_cell(row.cells[0], '  ' + name,   size=10, fill=fill)
    set_cell(row.cells[1], detail, italic=True, size=10, fill=fill, color_hex='444444')
    set_cell(row.cells[2], amount, size=10, align='right', fill=fill)

def fmt(n): return f'{int(n):,}'

# ── Update existing rows ──────────────────────────────────────────────────────
# R02 — PI: 15% -> 20% FTE
update_item_row(t.rows[2],
    'Principal Investigator (PhD, Assoc. Prof.)',
    '45,000 THB/mo x 20% FTE x 18 mo', '162,000', 'EBF5EB')

# R03 — Co-PI: 10% -> 15% FTE
update_item_row(t.rows[3],
    'Co-Principal Investigator (PhD, Asst. Prof.)',
    '35,000 THB/mo x 15% FTE x 18 mo', '94,500', 'FFFFFF')

# R04 — AI/ML Dev 1: fix label to Master's (rate unchanged)
update_item_row(t.rows[4],
    "AI/ML Developer 1 (Master's degree)",
    '22,000 THB/mo x 100% FTE x 18 mo', '396,000', 'EBF5EB')

# R05 — AI/ML Dev 2: fix label to Master's (rate unchanged)
update_item_row(t.rows[5],
    "AI/ML Developer 2 (Master's degree)",
    '22,000 THB/mo x 100% FTE x 18 mo', '396,000', 'FFFFFF')

# R06 — RA 1: revert to Bachelor's 18,000
update_item_row(t.rows[6],
    "Clinical Research Assistant 1 (Bachelor's degree)",
    '18,000 THB/mo x 100% FTE x 18 mo', '324,000', 'EBF5EB')

# ── Insert RA 2 row after R06 (before subtotal at R07) ───────────────────────
new_tr = copy.deepcopy(t.rows[6]._tr)
t.rows[6]._tr.addnext(new_tr)
# RA 2 is now R07; old subtotal shifted to R08
update_item_row(t.rows[7],
    "Clinical Research Assistant 2 (Bachelor's degree)",
    '18,000 THB/mo x 100% FTE x 18 mo', '324,000', 'FFFFFF')

# ── Update Cat 1 subtotal (now R08) ──────────────────────────────────────────
CAT1 = 162_000 + 94_500 + 396_000 + 396_000 + 324_000 + 324_000   # 1,696,500
set_cell(t.rows[8].cells[0], '  Subtotal — Category 1', bold=True, size=10, fill='D5F5E3')
set_cell(t.rows[8].cells[1], '', size=10, fill='D5F5E3')
set_cell(t.rows[8].cells[2], fmt(CAT1), bold=True, size=10, align='right', fill='D5F5E3')

# ── Recalculate downstream ────────────────────────────────────────────────────
CAT2   = 500_000
CAT3   = 2_310_000
CAT4   = 400_000
CAT5   = round((CAT1 + CAT2 + CAT3) * 0.10)   # 450,650
GRAND  = CAT1 + CAT2 + CAT3 + CAT4 + CAT5

# Update Cat 5 item + subtotal + compliance + grand total by scanning
for ri, row in enumerate(t.rows):
    txt = row.cells[0].text.strip()

    if 'MFU institutional overhead' in txt:
        base = CAT1 + CAT2 + CAT3
        set_cell(row.cells[1],
                 f'10% x (Cat 1+2+3) = 10% x {fmt(base)} THB',
                 italic=True, size=10, fill='EBF5EB')
        set_cell(row.cells[2], fmt(CAT5), size=10, align='right', fill='EBF5EB')

    elif 'Subtotal — Category 5' in txt:
        set_cell(row.cells[2], fmt(CAT5), bold=True, size=10, align='right', fill='D5F5E3')

    elif 'HSRI compliance' in txt:
        non_eq  = GRAND - CAT4
        pct1    = CAT1 / non_eq * 100
        pct2    = CAT2 / non_eq * 100
        pct5    = CAT5 / (CAT1+CAT2+CAT3) * 100
        flag    = '<= 30% OK' if pct1 <= 30 else f'> 30% — EXCEEDS CAP (requires HSRI Director approval)'
        msg = (f'HSRI compliance  |  '
               f'Personnel (Cat 1): {fmt(CAT1)} / {fmt(non_eq)} = {pct1:.1f}% {flag}   '
               f'Overhead (Cat 2): {pct2:.1f}% <= 15%   '
               f'Fee (Cat 5): {pct5:.1f}% <= 10%')
        set_cell(row.cells[0], msg, italic=True, size=9, fill='FEF9E7')

    elif 'GRAND TOTAL' in txt:
        set_cell(row.cells[2], fmt(GRAND) + ' THB',
                 bold=True, size=11, align='right',
                 fill='0F6E56', color_hex='FFFFFF')

# ── Update cover table row 6 ─────────────────────────────────────────────────
cover = doc.tables[0].rows[6].cells[0]
cover.paragraphs[0].clear()
cover.paragraphs[0].add_run(
    f'Requested Budget: {fmt(GRAND)} THB  '
    f'(HSRI 5-category structure, FY2570)')

# ── Print summary ─────────────────────────────────────────────────────────────
non_eq = GRAND - CAT4
print(f'Category 1 (Personnel):  {fmt(CAT1)} THB  ({CAT1/non_eq*100:.1f}% of non-equip total)')
print(f'Category 2 (Overhead):   {fmt(CAT2)} THB  ({CAT2/non_eq*100:.1f}%)')
print(f'Category 3 (Operations): {fmt(CAT3)} THB')
print(f'Category 4 (Equipment):  {fmt(CAT4)} THB')
print(f'Category 5 (Fee):        {fmt(CAT5)} THB  ({CAT5/(CAT1+CAT2+CAT3)*100:.1f}% of Cat1+2+3)')
print(f'GRAND TOTAL:             {fmt(GRAND)} THB')
print()
print(f'Personnel cap (HSRI limit 30%): {CAT1/non_eq*100:.1f}%'
      + (' OK' if CAT1/non_eq <= 0.30 else ' --- EXCEEDS 30% --- requires HSRI Director approval'))

doc.save(TMP)
os.replace(TMP, DOCX)
print('Saved.')
