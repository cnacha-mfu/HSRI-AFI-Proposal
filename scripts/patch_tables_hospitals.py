# -*- coding: utf-8 -*-
"""Patch remaining table cells to change SHPHs → district hospitals for validation phase."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)

def rewrite_para(para, new_text):
    if not para.runs:
        para.add_run(new_text); return
    f = para.runs[0]
    bold, italic, size, underline = f.bold, f.italic, f.font.size, f.underline
    color = None
    try:
        if f.font.color and f.font.color.type:
            color = f.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(new_text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

def fix(ti, ri, ci, pi, new_text):
    cell = doc.tables[ti].rows[ri].cells[ci]
    para = cell.paragraphs[pi]
    print(f'  T{ti}R{ri}C{ci}P{pi}: {repr(para.text[:80])} →')
    print(f'       {repr(new_text[:80])}')
    rewrite_para(para, new_text)

# Table 3 R3C1P0 — Phase 2 activities
fix(3, 3, 1, 0,
    'Deploy at 3–5 CCRU-network district hospitals | '
    'Validate on 150+ real AFI cases with blood-confirmed diagnoses (PCR/serology) | '
    'Measure sensitivity/specificity and usability | '
    'Test offline and connected modes | '
    'Iterate model')

# Table 4 R2C1P0 — CCRU primary role (single paragraph with \n)
fix(4, 2, 1, 0,
    '[1] ML Training Dataset: deliver de-identified AFI database (>500 cases) with confirmed diagnoses, '
    'lab results, and outcomes for training/validation and Bayesian Prior calibration\n'
    '[2] Clinical Trial Support: prospective data source for Phase 2 field validation at district hospitals '
    '(where blood draws enable confirmed diagnosis); clinical expert review of model outputs; '
    'coordination of 150+ case collection\n'
    '[3] AMR and Clinical Expertise: design of antibiotic stewardship signal; '
    'field research assistants and travel costs (via MFU–MORU subcontract)')

# Table 4 R2C2P0 — CCRU contributions
fix(4, 2, 2, 0,
    'AFI retrospective dataset (>500, labelled) | Disease prevalence priors | AMR/AST data | '
    'District hospital field network | Clinical validation site | '
    'Expert ground-truth review | Oxford/Wellcome affiliation')

# Table 4 R3C1P0 — Partner SHPHs primary role (now future deployment, not validation)
fix(4, 3, 1, 0,
    'Future deployment sites (post-validation) | Real end-users (CHWs/nurses) | '
    'Community engagement for scale-up')

# Table 5 R4C1P0 — months 9–12
fix(5, 4, 1, 0,
    'Deploy at 2 district hospitals (soft field trial) | '
    'Collect 75 cases with blood-confirmed diagnoses | '
    'Measure sensitivity/specificity Round 1 | '
    'Go/No-go review')

# Table 5 R5C1P0 — months 13–15
fix(5, 5, 1, 0,
    'Expand to ≥3 district hospitals | '
    'Complete 150+ cases | '
    'Measure AMR stewardship outcomes | '
    'Final clinical validation')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nDone — saved.')
