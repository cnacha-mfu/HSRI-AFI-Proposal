# -*- coding: utf-8 -*-
"""
Translate remaining English text in the budget table of Research_Proposal_TH.docx.
Uses hardcoded translations for known structural text + Google Translate for items.
"""
import sys, os, time, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from deep_translator import GoogleTranslator

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_TH.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_TH_tmp.docx')

doc = Document(DOCX)
t   = doc.tables[6]
tr  = GoogleTranslator(source='en', target='th')

# ── Hardcoded structural translations ──────────────────────────────────────────
FIXED = {
    # Category headers
    'Category 1 — Personnel Compensation (หมวดค่าตอบแทน)':
        'หมวดที่ 1 — ค่าตอบแทนบุคลากร (หมวดค่าตอบแทน)',
    'Category 2 — Project Management Overhead (หมวดค่าบริหารจัดการ)':
        'หมวดที่ 2 — ค่าบริหารจัดการโครงการ (หมวดค่าบริหารจัดการ)',
    'Category 3 — Research Operations (หมวดค่าดำเนินงาน)':
        'หมวดที่ 3 — ค่าดำเนินงานวิจัย (หมวดค่าดำเนินงาน)',
    'Category 4 — Equipment (หมวดค่าครุภัณฑ์)':
        'หมวดที่ 4 — ค่าครุภัณฑ์ (หมวดค่าครุภัณฑ์)',
    'Category 5 — Institutional Overhead Fee (หมวดค่าธรรมเนียมหน่วยงานผู้รับทุน)':
        'หมวดที่ 5 — ค่าธรรมเนียมสถาบัน (หมวดค่าธรรมเนียมหน่วยงานผู้รับทุน)',
    # Compliance note header items
    'Subtotal — Category 1': 'รวมย่อย — หมวดที่ 1',
    'Subtotal — Category 2': 'รวมย่อย — หมวดที่ 2',
    'Subtotal — Category 3': 'รวมย่อย — หมวดที่ 3',
    'Subtotal — Category 4': 'รวมย่อย — หมวดที่ 4',
    'Subtotal — Category 5': 'รวมย่อย — หมวดที่ 5',
    # Personnel
    'Principal Investigator (PhD)':
        'หัวหน้าโครงการวิจัย (ปริญญาเอก, รศ.)',
    'Co-Principal Investigator (PhD, Asst. Prof.)':
        'ผู้ร่วมวิจัยหลัก (ปริญญาเอก, ผศ.)',
    "Clinical Research Assistant (Bachelor's degree)":
        'ผู้ช่วยวิจัยทางคลินิก (ปริญญาตรี)',
    # CCRU subcontract
    'CCRU/MORU subcontract — Field RA salary  [MFU-MORU subcontract agreement]':
        'สัญญาช่วง CCRU/MORU — เงินเดือนผู้ช่วยวิจัยภาคสนาม [สัญญาช่วง MFU-MORU]',
    # Cat 3 items
    'Field data collection travel — MFU research team':
        'ค่าเดินทางเก็บข้อมูลภาคสนาม — ทีมวิจัย MFU',
    'CHW and nurse training workshops':
        'การอบรมเชิงปฏิบัติการสำหรับ อสม. และพยาบาล',
    'Participant engagement and consent process costs':
        'ค่าดำเนินการมีส่วนร่วมผู้เข้าร่วมและกระบวนการรับทราบยินยอม',
    'Data annotation and labeling services':
        'บริการถอดความและติดป้ายกำกับข้อมูล',
    'Clinical expert advisory — 3 external reviewers + UAT facilitation':
        'ที่ปรึกษาผู้เชี่ยวชาญทางคลินิก — ผู้ทบทวนภายนอก 3 คน + ผู้ดำเนิน UAT',
    'Software tool subscriptions (LLM API, annotation platform, NLP tools, CI/CD)':
        'ค่าสมัครสมาชิกซอฟต์แวร์ (LLM API, แพลตฟอร์ม Annotation, เครื่องมือ NLP, CI/CD)',
    # Cat 4
    'Digital vital-signs monitors (SpO2, BP, temperature) — field validation':
        'เครื่องวัดสัญญาณชีพดิจิทัล (SpO2, ความดัน, อุณหภูมิ) — การทดสอบภาคสนาม',
    # Cat 5
    'MFU institutional overhead — utilities, workspace, admin, shared IT infrastructure':
        'ค่าธรรมเนียมสถาบัน MFU — สาธารณูปโภค พื้นที่ทำงาน ธุรการ โครงสร้างพื้นฐาน IT',
    # Calc bases that weren't translated
    '3 reports x lump sum': '3 รายงาน x เหมาจ่าย',
    '~52,778 THB/mo average × 18 months (consolidated)':
        'เฉลี่ย ~52,778 บาท/เดือน × 18 เดือน (รวมเบ็ดเสร็จ)',
    '20 planned field visits × 10,000 THB avg. (travel + accommodation)':
        '20 ครั้งภาคสนาม × 10,000 บาท เฉลี่ย (เดินทาง + ที่พัก)',
    '5 training sites × 24,000 THB per site (lump sum)':
        '5 พื้นที่ฝึกอบรม × 24,000 บาท/พื้นที่ (เหมาจ่าย)',
    '5 sites × 16,000 THB per site (consent facilitation + participant incentives, lump sum)':
        '5 พื้นที่ × 16,000 บาท/พื้นที่ (ดำเนินการยินยอม + สิ่งจูงใจผู้เข้าร่วม, เหมาจ่าย)',
    'CDSS rule validation, UAT protocol review, 2 rounds clinical feedback':
        'ตรวจสอบกฎ CDSS, ทบทวนโปรโตคอล UAT, ข้อเสนอแนะทางคลินิก 2 รอบ',
    '22,000 THB/mo x 100% FTE x 18 mo': '22,000 บาท/เดือน × 100% FTE × 18 เดือน',
    '18,000 THB/mo x 100% FTE x 18 mo': '18,000 บาท/เดือน × 100% FTE × 18 เดือน',
    '35,000 THB/mo x 19% FTE x 18 mo': '35,000 บาท/เดือน × 19% FTE × 18 เดือน',
    '25,000 บาท/เดือน x 12 เดือน': '25,000 บาท/เดือน × 12 เดือน',
    # Compliance note
    '21.0% of total excl. equipment — compliant with <=30% HSRI cap':
        '21.0% ของยอดรวมไม่รวมครุภัณฑ์ — เป็นไปตามเพดาน สวรส. ไม่เกิน 30%',
    '9.9% of total excl. equipment — compliant with <=15% HSRI cap':
        '9.9% ของยอดรวมไม่รวมครุภัณฑ์ — เป็นไปตามเพดาน สวรส. ไม่เกิน 15%',
    'Includes MFU-MORU subcontract (HSRI Sec. 3.2.3)':
        'รวมสัญญาช่วง MFU-MORU (ตามระเบียบ สวรส. ข้อ 3.2.3)',
    'All items >= 20,000 THB/unit; excluded from Cat 1 percentage denominator':
        'ทุกรายการ >= 20,000 บาท/หน่วย; ไม่นำมาคำนวณสัดส่วนบุคลากร',
    '10% of (Cat 1+2+3) = 4,152,200 THB — compliant with <=10% HSRI cap':
        '10% ของ (หมวด 1+2+3) = 4,152,200 บาท — เป็นไปตามเพดาน สวรส. ไม่เกิน 10%',
}

def is_mainly_thai(text):
    thai = sum(1 for c in text if '\u0e00' <= c <= '\u0e7f')
    return thai / max(len(text), 1) > 0.5

def is_number_only(text):
    return bool(re.match(r'^[\d,\.]+\s*(THB|บาท)?$', text.strip()))

def google_translate(text):
    for attempt in range(3):
        try:
            result = tr.translate(text)
            time.sleep(0.3)
            return result if result else text
        except Exception as e:
            if attempt == 2:
                print(f'  FAIL: {text[:40]!r} — {e}')
                return text
            time.sleep(3)
    return text

def translate_text(text):
    text_s = text.strip()
    if not text_s or is_number_only(text_s) or is_mainly_thai(text_s):
        return text
    # Check hardcoded first
    if text_s in FIXED:
        return FIXED[text_s]
    # Also check with leading spaces stripped
    for k, v in FIXED.items():
        if text_s.strip() == k.strip():
            return v
    # Fall back to Google Translate
    return google_translate(text_s)

def apply_to_para(para, new_text):
    """Replace paragraph text preserving formatting of first run."""
    if not para.runs:
        para.add_run(new_text)
        return
    fmt = para.runs[0]
    bold, italic, size, underline = fmt.bold, fmt.italic, fmt.font.size, fmt.underline
    color = None
    try:
        if fmt.font.color and fmt.font.color.type:
            color = fmt.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(new_text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

# ── Process budget table ───────────────────────────────────────────────────────
seen_cells = set()
for ri, row in enumerate(t.rows):
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen_cells:
            continue
        seen_cells.add(cid)
        for pi, para in enumerate(cell.paragraphs):
            orig = para.text.strip()
            if not orig:
                continue
            thai = translate_text(orig)
            if thai != orig:
                print(f'R{ri:02d}P{pi}: {orig[:45]!r} → {thai[:45]!r}')
                apply_to_para(para, thai)

# ── Grand total row: also fix "THB" → "บาท" ──────────────────────────────────
for row in t.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            txt = para.text.strip()
            if 'GRAND TOTAL' in txt or '4,967,420 THB' in txt:
                new = txt.replace('GRAND TOTAL', 'ยอดรวมทั้งหมด').replace('Research Framework', 'กรอบการวิจัย').replace(' THB', ' บาท')
                if new != txt:
                    print(f'Grand total fix: {new[:70]}')
                    apply_to_para(para, new)

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nDone — saved.')
