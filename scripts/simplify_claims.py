# -*- coding: utf-8 -*-
"""
Restructure Cat 3 for minimum documentation burden.
No amount changes — grand total stays 4,967,420 THB.

Changes:
  1. Merge Cloud AI training + inference → single 950,000 line
  2. CCRU field logistics: 150,000 → 200,000 (absorbs comm allowances)
  3. Field travel: "40 visits x 5,000" → "20 visits x 10,000" (half the claims)
  4. CHW training: "5 sites x 3 sessions x 8,000" → "5 sites x 24,000 lump sum"
  5. Participant engagement: "400 participants" → "5 sites x 16,000 lump sum"
  6. Remove comm allowances row (50,000 absorbed into CCRU logistics above)
  7. AI dev licenses: USD → THB denomination (same amount 126,000)
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)
t   = doc.tables[6]

def fmt(n): return f'{int(n):,}'

def shade(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex); tcPr.append(shd)

def set_cell(cell, text, bold=False, italic=False, size=10,
             align='left', fill=None, color_hex=None):
    if fill: shade(cell, fill)
    para = cell.paragraphs[0]; para.clear()
    para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT  if align == 'right'  else
                      WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else
                      WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))

# ── Step 1: Targeted row edits before any deletions ───────────────────────────
for row in t.rows:
    c0 = row.cells[0].text.strip()

    # Merge Cloud AI: update training row to combined 950K
    if 'Cloud AI — model' in c0 or 'Cloud AI — model Deployment' in c0:
        set_cell(row.cells[0], '  Cloud AI infrastructure — AWS/GCP (compute, LLM API, storage, CI/CD)', size=10)
        set_cell(row.cells[1], '~52,778 THB/mo average × 18 months (consolidated)', italic=True, size=10, color_hex='444444')
        set_cell(row.cells[2], fmt(950_000), size=10, align='right')
        print('Updated: Cloud AI → 950,000 (merged)')

    # CCRU field logistics: 150K → 200K
    elif 'CCRU/MORU subcontract — Field travel' in c0:
        set_cell(row.cells[1], 'Lump sum; 5 SHPH sites across Chiang Rai province (incl. CHW field comm.)', italic=True, size=10, color_hex='444444')
        set_cell(row.cells[2], fmt(200_000), size=10, align='right')
        print('Updated: CCRU field logistics → 200,000')

    # Field travel: 40 visits → 20 visits
    elif 'Field data collection travel' in c0:
        set_cell(row.cells[1], '20 planned field visits × 10,000 THB avg. (travel + accommodation)', italic=True, size=10, color_hex='444444')
        print('Updated: field travel basis → 20 visits × 10,000')

    # CHW training: per-session → per-site lump sum
    elif 'CHW and nurse training' in c0:
        set_cell(row.cells[1], '5 training sites × 24,000 THB per site (lump sum)', italic=True, size=10, color_hex='444444')
        print('Updated: CHW training → 5 sites × 24,000')

    # Participant engagement: 400 persons → per-site lump sum
    elif 'Participant engagement' in c0:
        set_cell(row.cells[1], '5 sites × 16,000 THB per site (consent facilitation + participant incentives, lump sum)', italic=True, size=10, color_hex='444444')
        print('Updated: participant engagement → 5 sites × 16,000')

    # AI dev licenses: USD → THB
    elif 'AI development tool licenses' in c0:
        set_cell(row.cells[1], '3,500 THB/developer/mo × 2 developers × 18 mo', italic=True, size=10, color_hex='444444')
        print('Updated: AI dev licenses → THB denomination')

# ── Step 2: Remove Cloud AI inference row and comm allowances row ──────────────
to_remove = []
for row in t.rows:
    c0 = row.cells[0].text.strip()
    if 'Cloud AI — inference' in c0:
        to_remove.append(('Cloud AI inference row', row._tr))
    elif 'Field communication allowances' in c0:
        to_remove.append(('Comm allowances row', row._tr))

for label, tr in to_remove:
    t._tbl.remove(tr)
    print(f'Removed: {label}')

# ── Step 3: Re-stripe Cat 3 item rows ────────────────────────────────────────
# Collect Cat 3 items between category header and subtotal
in_cat3 = False
cat3_item_rows = []
for row in t.rows:
    c0 = row.cells[0].text.strip()
    if 'Category 3' in c0 and 'Research Operations' in c0:
        in_cat3 = True
        continue
    if in_cat3:
        if 'Subtotal — Category 3' in c0:
            break
        cat3_item_rows.append(row)

for ii, row in enumerate(cat3_item_rows):
    fill = 'EBF5EB' if ii % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        shade(cell, fill)
print(f'Re-striped {len(cat3_item_rows)} Cat 3 item rows')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nSaved. Grand total unchanged: 4,967,420 THB')
print('\nDocumentation simplification summary:')
print('  Cloud AI: 2 separate AWS/GCP lines → 1 consolidated line')
print('  Field travel: 40 claims → 20 claims')
print('  CHW training: 15 session reports → 5 site-level lump sums')
print('  Participant engagement: 400 signatures → 5 site-level lump sums')
print('  Comm allowances: 500 SIM records → removed (in CCRU subcontract)')
print('  AI dev licenses: USD forex docs → THB, no currency conversion needed')
