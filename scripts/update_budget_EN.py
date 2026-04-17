# -*- coding: utf-8 -*-
"""
Replace Table 7 (budget) in Research_Proposal_EN.docx with HSRI-compliant
5-category detailed budget and update the cover-table budget cell.

HSRI rules applied:
  Cat 1 (Personnel)        : <= 30% of total excl. equipment
  Cat 2 (Overhead/Mgmt)    : <= 15% of total excl. equipment
  Cat 5 (Institutional fee): <= 10% of (Cat 1 + Cat 2 + Cat 3)
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)

# ── Budget structure (HSRI 5-category framework) ─────────────────────────────
#
# Grand total target: exactly 5,000,000 THB
#   Cat 1+2+3 = 4,182,500  →  institutional fee = 9.98% (< 10 %)
#   Personnel  = 1,372,500 / (5,000,000 – 400,000) = 29.8%  (< 30 %)
#   Overhead   =   500,000 / 4,600,000              = 10.9%  (< 15 %)
# ─────────────────────────────────────────────────────────────────────────────

BUDGET = [
    {
        'cat'  : 'Category 1 — Personnel Compensation (หมวดค่าตอบแทน)',
        'color': '1B4F72',
        'note' : '29.8% of total excl. equipment — compliant with <= 30% HSRI cap',
        'items': [
            ('Principal Investigator — PhD, Assoc. Prof.',
             '45,000 THB/mo x 20% FTE x 18 mo', 162_000),
            ('Co-Principal Investigator — PhD, Asst. Prof.',
             '35,000 THB/mo x 15% FTE x 18 mo', 94_500),
            ('AI/ML Developer 1 — Master\'s degree',
             '22,000 THB/mo x 100% FTE x 18 mo', 396_000),
            ('AI/ML Developer 2 — Master\'s degree',
             '22,000 THB/mo x 100% FTE x 18 mo', 396_000),
            ('Clinical Research Assistant — Bachelor\'s degree',
             '18,000 THB/mo x 100% FTE x 18 mo', 324_000),
        ],
    },
    {
        'cat'  : 'Category 2 — Project Management Overhead (หมวดค่าบริหารจัดการ)',
        'color': '145A32',
        'note' : '10.9% of total excl. equipment — compliant with <= 15% HSRI cap',
        'items': [
            ('Documentation — inception, interim and final research reports',
             '3 reports x lump sum', 80_000),
            ('Office supplies and stationery',
             'Lump sum over 18 months', 36_000),
            ('Telephone and internet — project coordination charges',
             '2,000 THB/mo x 12 mo', 24_000),
            ('Liaison travel — HSRI / MoPH progress review meetings',
             '4 trips x 15,000 THB (Bangkok)', 60_000),
            ('Printing and reproduction — reports, manuals, consent forms',
             'Lump sum', 30_000),
            ('Publication fees',
             '2 peer-reviewed journal articles x 60,000 THB', 120_000),
            ('Conference presentations',
             '2 conferences (1 domestic + 1 regional) x 50,000 THB', 100_000),
            ('Policy dissemination workshop — Chiang Rai',
             '1 workshop, ~50 participants', 50_000),
        ],
    },
    {
        'cat'  : 'Category 3 — Research Operations (หมวดค่าดำเนินงาน)',
        'color': '4A235A',
        'note' : 'Includes MFU-MORU subcontract (budgeted per HSRI Sec. 3.2.3)',
        'items': [
            ('Cloud AI infrastructure — AWS/GCP (GPU compute, LLM API, storage, CI/CD)',
             '50,000 THB/mo x 18 mo', 900_000),
            ('CCRU/MORU subcontract — Field RA salary  [MFU-MORU subcontract agreement]',
             '25,000 THB/mo x 12 mo', 300_000),
            ('CCRU/MORU subcontract — Field travel and site logistics',
             'Lump sum; 5 SHPH sites across Chiang Rai province', 150_000),
            ('Field data collection travel — MFU research team',
             'Est. 40 site visits x 5,000 THB avg.', 200_000),
            ('CHW and nurse training workshops',
             '5 sites x 3 sessions x 8,000 THB per session', 120_000),
            ('Participant engagement and consent process costs',
             'Incentives + refreshments; 400 enrolled participants', 80_000),
            ('Data annotation and labeling services',
             'Audio transcription + clinical entity labeling (external service)', 150_000),
            ('External statistical analysis consultant',
             '1 consultant; fixed-fee contract', 80_000),
            ('Clinical expert advisory — 3 external reviewers',
             'CDSS rule validation + UAT protocol review', 80_000),
            ('Software tool subscriptions (LLM API credits, annotation platform, NLP tools)',
             'Multiple subscriptions < 20,000 THB/unit; classified as operations per HSRI', 150_000),
            ('Field communication allowances — SIM/data for CHW tablets',
             '100 CHW devices x 1,000 THB/device x 10 mo', 100_000),
        ],
    },
    {
        'cat'  : 'Category 4 — Equipment (หมวดค่าครุภัณฑ์)',
        'color': '7D6608',
        'note' : 'All items >= 20,000 THB/unit; excluded from Cat 1 percentage denominator',
        'items': [
            ('GPU development workstation — on-site model training and inference',
             '1 unit x 80,000 THB', 80_000),
            ('High-performance research laptops',
             '3 units x 35,000 THB', 105_000),
            ('Field server / NAS data storage — offline SHPH deployment',
             '1 unit x 60,000 THB', 60_000),
            ('Network infrastructure — managed switch + wireless router (SHPH)',
             '1 set x 45,000 THB', 45_000),
            ('Digital vital-signs monitors (SpO2, BP, temperature) — field validation',
             '5 sets x 22,000 THB', 110_000),
        ],
    },
    {
        'cat'  : 'Category 5 — Institutional Overhead Fee (หมวดค่าธรรมเนียมหน่วยงานผู้รับทุน)',
        'color': '2E4053',
        'note' : '9.98% of (Cat 1 + 2 + 3) = 4,182,500 THB — compliant with <= 10% HSRI cap',
        'items': [
            ('MFU institutional overhead — utilities, shared workspace, admin, shared IT infrastructure',
             '9.98% x (Category 1 + 2 + 3) = 9.98% x 4,182,500 THB', 417_500),
        ],
    },
]

for cat in BUDGET:
    cat['subtotal'] = sum(item[2] for item in cat['items'])

grand_total = sum(cat['subtotal'] for cat in BUDGET)

def fmt(n):
    return f'{int(n):,}'

# ── Cell helpers ──────────────────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, italic=False, size=10,
             color_hex=None, align='left', fill=None):
    if fill:
        shade_cell(cell, fill)
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = {
        'right' : WD_ALIGN_PARAGRAPH.RIGHT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

def add_note_para(cell, text):
    """Add a smaller italic note paragraph inside a cell."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── Count rows needed ─────────────────────────────────────────────────────────
# 1 header + per category: 1 cat-header + n items + 1 subtotal + 1 compliance + 1 grand total
n_rows = 1
for cat in BUDGET:
    n_rows += 1 + len(cat['items']) + 1
n_rows += 2   # compliance note + grand total

# ── Build replacement table ───────────────────────────────────────────────────
tbl = doc.add_table(rows=n_rows, cols=3)
tbl.style = 'Table Grid'

COL_W = [Cm(7.5), Cm(6.5), Cm(3.5)]
ri = 0

# ── Table header row ──────────────────────────────────────────────────────────
hdr = tbl.rows[ri]
set_cell(hdr.cells[0], 'Budget Category / Line Item',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56')
set_cell(hdr.cells[1], 'Calculation Basis / Details',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56')
set_cell(hdr.cells[2], 'Amount (THB)',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56', align='center')
ri += 1

# ── Category blocks ───────────────────────────────────────────────────────────
for cat_idx, cat in enumerate(BUDGET):

    # Category header — spans all 3 columns
    cat_row = tbl.rows[ri]
    merged = cat_row.cells[0].merge(cat_row.cells[2])
    set_cell(merged, cat['cat'],
             bold=True, size=10.5, color_hex='FFFFFF', fill=cat['color'])
    if 'note' in cat:
        add_note_para(merged, cat['note'])
    ri += 1

    # Item rows
    for ii, (name, detail, amount) in enumerate(cat['items']):
        row = tbl.rows[ri]
        fill = 'EBF5EB' if ii % 2 == 0 else 'FFFFFF'
        for ci in range(3):
            shade_cell(row.cells[ci], fill)
        set_cell(row.cells[0], '  ' + name, size=10)
        set_cell(row.cells[1], detail, size=10, italic=True, color_hex='444444')
        set_cell(row.cells[2], fmt(amount), size=10, align='right')
        ri += 1

    # Subtotal row
    sub = tbl.rows[ri]
    for ci in range(3):
        shade_cell(sub.cells[ci], 'D5F5E3')
    set_cell(sub.cells[0],
             f'  Subtotal — Category {cat_idx + 1}',
             bold=True, size=10)
    set_cell(sub.cells[1], '', size=10)
    set_cell(sub.cells[2], fmt(cat['subtotal']),
             bold=True, size=10, align='right')
    ri += 1

# ── Compliance summary row ────────────────────────────────────────────────────
comp_row = tbl.rows[ri]
merged_comp = comp_row.cells[0].merge(comp_row.cells[2])
shade_cell(comp_row.cells[0], 'FEF9E7')
set_cell(comp_row.cells[0],
         'HSRI compliance  |  '
         'Personnel (Cat 1): 1,372,500 / 4,600,000 = 29.8% <= 30%   '
         'Overhead (Cat 2): 500,000 / 4,600,000 = 10.9% <= 15%   '
         'Institutional fee (Cat 5): 417,500 / 4,182,500 = 9.98% <= 10%',
         size=9, italic=True, color_hex='555555')
ri += 1

# ── Grand total row ───────────────────────────────────────────────────────────
gt_row = tbl.rows[ri]
gt_merged = gt_row.cells[0].merge(gt_row.cells[1])
shade_cell(gt_row.cells[0], '0F6E56')
shade_cell(gt_row.cells[2], '0F6E56')
set_cell(gt_row.cells[0],
         'GRAND TOTAL  (HSRI FY2570 — Research Framework 3.1.1 + 3.1.2)',
         bold=True, size=11, color_hex='FFFFFF')
set_cell(gt_row.cells[2],
         fmt(grand_total) + ' THB',
         bold=True, size=11, color_hex='FFFFFF', align='right')

# ── Set column widths on unmerged rows ────────────────────────────────────────
for row in tbl.rows:
    cells = row.cells
    distinct = list(dict.fromkeys(id(c) for c in cells))  # preserve order, deduplicate
    if len(distinct) == 3:
        for ci, w in enumerate(COL_W):
            cells[ci].width = w

# ── Swap old table (index 7) for new one ─────────────────────────────────────
old_tbl_elem = doc.tables[7]._tbl
new_tbl_elem = tbl._tbl
body = doc.element.body
body.remove(new_tbl_elem)                       # remove from where add_table appended it
idx = list(body).index(old_tbl_elem)
body.remove(old_tbl_elem)
body.insert(idx, new_tbl_elem)
print('Budget table replaced at body position', idx)

# ── Update cover table (Table 0, row 6) ──────────────────────────────────────
cover_cell = doc.tables[0].rows[6].cells[0]
for run in cover_cell.paragraphs[0].runs:
    run.text = ''
if cover_cell.paragraphs[0].runs:
    cover_cell.paragraphs[0].runs[0].text = (
        'Requested Budget: 5,000,000 THB  '
        '(Cat 1 Personnel 29.8% | Cat 2 Overhead 10.9% | Cat 5 Fee 9.98% — '
        'HSRI 5-category structure, FY2570)'
    )
else:
    cover_cell.paragraphs[0].add_run(
        'Requested Budget: 5,000,000 THB  '
        '(Cat 1 Personnel 29.8% | Cat 2 Overhead 10.9% | Cat 5 Fee 9.98% — '
        'HSRI 5-category structure, FY2570)'
    )
print('Cover table budget cell updated.')

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(TMP)
os.replace(TMP, DOCX)
print(f'Saved. Grand total: {fmt(grand_total)} THB')
print()
for i, cat in enumerate(BUDGET, 1):
    pct_nonequip = cat['subtotal'] / (grand_total - BUDGET[3]['subtotal']) * 100
    print(f'  Cat {i}: {fmt(cat["subtotal"]):>12} THB  ({pct_nonequip:.1f}% of non-equip total)')
print(f'  {"TOTAL":>4}: {fmt(grand_total):>12} THB')
