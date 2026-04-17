# -*- coding: utf-8 -*-
"""Create English version of the AFI AI Screening research proposal."""
import sys, os, zipfile, shutil
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
OUT_FILE = os.path.join(FOLDER, 'Research_Proposal_EN.docx')

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x53, 0x4A, 0xB7)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

# ═════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RESEARCH PROPOSAL')
r.bold = True; r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Health Systems Research Institute (HSRI) — Fiscal Year 2027')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI-Assisted Screening for Acute Febrile Illness\nin Remote Areas, Chiang Rai, Northern Thailand')
r.bold = True; r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x0F, 0x6E, 0x56)

doc.add_paragraph()

# ── Cover table ───────────────────────────────────────────────────────────────
tbl = doc.add_table(rows=9, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data = [
    ('Lead Institution',
     'Mae Fah Luang University (MFU) — School of Applied Digital Technology (ADT)'),
    ('Partner Institution',
     'Chiang Rai Clinical Research Unit (CCRU) / Mahidol Oxford Tropical Medicine Research Unit (MORU)'),
    ('Principal Investigator',
     'Assoc. Prof. Dr. Nacha Choldarongkul (Dean, School of ADT)\nnacha.cho@mfu.ac.th  |  0 5391 6744'),
    ('Co-Investigators',
     'Asst. Prof. Dr. Phattaramon Wuttipitayamongkol (MFU ADT, Co-PI)\nDr. Carlo Perrone (CCRU / MORU, Chiang Rai)'),
    ('HSRI Research Framework',
     '3.1.1 — Improving health service efficiency through medical technology\n'
     '3.1.2 — Antimicrobial resistance (AMR) management research'),
    ('Project Duration',
     '18 months (October 2026 – March 2028)  |  FY 2027–2028'),
    ('Requested Budget',
     '5,005,000 THB (approx.) — incl. 10% institutional overhead and CCRU field support'),
    ('Study Sites',
     'Chiang Rai: CCRU/MORU + 3–5 partner Sub-district Health Promotion Hospitals (SHPH) + MFU'),
    ('Technology Readiness Level',
     'TRL 3 (proof-of-concept) to TRL 6 (validated field prototype)'),
]

for i, (label, value) in enumerate(rows_data):
    row = tbl.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
    row.cells[0].paragraphs[0].runs[0].bold = True
    shade_cell(row.cells[0], 'E8F5F0')

set_col_widths(tbl, [4.5, 12.0])
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 1: Background and Problem Statement')

h2('1.1  Problem Situation')
body(
    'Acute Febrile Illness (AFI) is a leading cause of preventable death in rural communities '
    'of northern Thailand, particularly among highland ethnic minority populations with limited '
    'access to health services. The principal diseases encountered in this area include scrub '
    'typhus, dengue haemorrhagic fever, leptospirosis, and bacteraemia/sepsis.'
)
body(
    'A critical gap at the primary care level is that community health workers (CHW/อสม.) and '
    'nurses at Sub-district Health Promotion Hospitals (SHPH/รพ.สต.) lack a practical decision-'
    'support tool to determine which febrile patients require urgent referral and which can be '
    'safely managed locally. This diagnostic uncertainty also drives inappropriate antibiotic '
    'prescribing, contributing to antimicrobial resistance (AMR) patterns documented by CCRU '
    'in Chiang Rai province.'
)

h2('1.2  Knowledge and Innovation Gaps')
body('A review of the literature and the Thai health system landscape reveals the following:')
bullet(
    'No Thai-language clinical decision support system (CDSS) for AFI currently exists that is '
    'designed for resource-limited settings.'
)
bullet(
    'Existing health AI in Thailand focuses predominantly on medical imaging (e.g., Chest 4 All AI) '
    'rather than differential diagnosis for infectious diseases.'
)
bullet(
    'CCRU/MORU holds the most comprehensive AFI clinical and AMR dataset in the region — covering '
    'more than 500 retrospective cases from Chiang Rai and northern Thailand with confirmed '
    'diagnoses, laboratory results, epidemiological data, and treatment outcomes — forming an '
    'invaluable training dataset and Bayesian Prior source for ML model development.'
)
bullet(
    'Mae Fah Luang University (MFU) has demonstrated expertise in Thai-language AI and NLP.'
)

h2('1.3  Alignment with HSRI Research Framework FY 2027')
tbl2 = doc.add_table(rows=6, cols=2)
tbl2.style = 'Table Grid'
align_data = [
    ('HSRI Framework', 'Alignment'),
    ('3.1.1 — Improve efficiency through medical technology',
     'Develop an AI platform for febrile illness diagnosis and risk prediction; '
     'increase access to services in resource-limited settings'),
    ('3.1.2 — AMR management research',
     'Embedded antibiotic stewardship signal reduces unnecessary antibiotic '
     'prescribing in AFI cases; aligned with National AMR Plan 2023–2027'),
    ('KR1 — Reduce infectious disease burden',
     'AI system reduces late diagnosis of scrub typhus and sepsis'),
    ('KR2 — Multi-sector collaboration network',
     '1 network: MFU (university) + CCRU/MORU (international research) + '
     '≥3 SHPHs (service) + Chiang Rai PHO (government)'),
    ('KR3 / KR5 — Technology deployed; population served',
     '1 AI AFI screening tool deployed in ≥3 SHPHs; ≥150 validated cases; '
     'catchment population ≥10,000 reached in scale-up phase'),
]
for i, (a, b) in enumerate(align_data):
    tbl2.rows[i].cells[0].text = a
    tbl2.rows[i].cells[1].text = b
    for cell in tbl2.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10.5)
    if i == 0:
        for cell in tbl2.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '0F6E56')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        shade_cell(tbl2.rows[i].cells[0], 'E8F5F0')
set_col_widths(tbl2, [5.5, 11.0])
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 2: Objectives and Targets')

h2('2.1  Primary Objective')
body(
    'To develop and evaluate the effectiveness of an AI-powered clinical decision support system '
    '(AI-CDSS) for screening patients presenting with acute febrile illness at primary care level '
    'in remote areas of Chiang Rai, operable by community health workers in the field.'
)

h2('2.2  Secondary Objectives')
bullet(
    'Develop a Thai-language Clinical NLP module to extract symptoms, exposure history, and vital '
    'signs from voice recordings or structured text input.'
)
bullet(
    'Design and validate an antibiotic stewardship signal embedded in the system to reduce '
    'unnecessary antibiotic prescribing in AFI cases.'
)
bullet(
    'Evaluate the usability and clinical accuracy of the system under real-world conditions at '
    'CCRU-network SHPHs.'
)
bullet(
    'Build the system on an Offline-first architecture using on-device SQLite storage, with '
    'automatic FHIR R4 sync to HIS when internet connectivity is available — ensuring full '
    'functionality in remote areas without internet, and preparing for integration with the '
    'Ministry of Public Health HIS/EMR infrastructure in a subsequent phase.'
)

h2('2.3  Key Performance Indicators (KPIs)')
tbl3 = doc.add_table(rows=7, cols=3)
tbl3.style = 'Table Grid'
kpi_data = [
    ('KPI', 'Target', 'Measurement Method'),
    ('Sensitivity (AFI screening)', '≥ 85%', 'Compared against specialist physician diagnosis'),
    ('Specificity (avoid over-referral)', '≥ 75%', 'Compared against specialist physician diagnosis'),
    ('Antibiotic appropriateness rate', 'Increase ≥ 15%', 'Before-vs-after system deployment'),
    ('Usability score', '≥ 70 / 100', 'CHW and nurse questionnaire'),
    ('Word Error Rate (medical terms)', '≤ 10%', 'Tested on medical speech test set'),
    ('Validated patient cases', '≥ 150 cases', 'From ≥3 SHPH partner sites'),
]
for i, row_data in enumerate(kpi_data):
    for j, val in enumerate(row_data):
        tbl3.rows[i].cells[j].text = val
        tbl3.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
    if i == 0:
        for cell in tbl3.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '534AB7')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif i % 2 == 0:
        for cell in tbl3.rows[i].cells:
            shade_cell(cell, 'EEEDFE')
set_col_widths(tbl3, [6.5, 3.5, 6.5])
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 3: Conceptual Framework and Research Methodology')

h2('3.1  Conceptual Framework')
body(
    'This project applies a Human-centred Design approach with community health workers (CHWs) '
    'and nurses at SHPHs as the primary user group. The system functions as a '
    '"Structured Checklist with Probabilistic Guidance" — a decision-support tool, not an '
    'automated diagnostic system. Clinicians and health personnel retain full responsibility '
    'for all final decisions.'
)
body(
    'Because SHPHs in remote Chiang Rai may lack stable internet connectivity, the system is '
    'built on an Offline-first principle: all core functions — voice input, text processing, '
    'risk assessment, and result recording — operate fully without any internet connection or '
    'HIS access. All data are stored in an on-device SQLite database and automatically synced '
    'to the HIS when connectivity is available. Three operating modes are supported: '
    '(1) Standalone — no internet; (2) Cached — intermittent connectivity; '
    '(3) Connected — live HIS sync.'
)

h2('System Architecture — 5 Processing Layers')
body(
    'Layer 1 — Input capture: The system accepts data from four sources: natural Thai-language '
    'speech from the patient or CHW; structured questionnaire responses; manually entered vital '
    'signs; and prior medical history retrieved from the HIS when connected, or manually entered '
    'in Offline mode. Speech is transcribed by a Whisper model fine-tuned for Thai medical '
    'vocabulary, targeting a Word Error Rate (WER) of no more than 10%.',
    indent=True
)
body(
    'Layer 2 — Structured clinical data extraction: Transcribed text is cleaned and normalised, '
    'then a zero-shot Named Entity Recognition (NER) model (GLiNER) extracts clinically relevant '
    'entities — including symptoms, exposure history, duration of illness, current medications, '
    'blood pressure, body temperature, and SpO2 — and populates them into a predefined AFI '
    'Clinical Schema.',
    indent=True
)
body(
    'Layer 3 — AI risk assessment: The structured clinical data are passed to a risk-scoring '
    'model that combines a Bayesian Prior derived from CCRU retrospective data (reflecting local '
    'disease prevalence in Chiang Rai) with an LLM reasoning layer for additional clinical '
    'inference. Outputs include probability scores for four target diseases — scrub typhus, '
    'dengue haemorrhagic fever, leptospirosis, and bacteraemia/sepsis — alongside a triage '
    'recommendation and an AMR Stewardship Signal.',
    indent=True
)
body(
    'Layer 4 — Review, record, and sync: All outputs are presented to a clinician or nurse for '
    'review, editing, and explicit approval before any record is saved. Upon approval, data are '
    'written to on-device SQLite storage immediately and automatically synced to the HIS via '
    'FHIR R4 when internet is available (optional). The system also dispatches automated '
    'follow-up alerts to the CHW on Day 3 and Day 7 post-assessment.',
    indent=True
)

body('Figure 3-1 and Figure 3-2 illustrate the system architecture and patient journey respectively.')

# Insert Figure 1
doc.add_paragraph()
img_path_1 = os.path.join(FOLDER, 'word', 'media', 'image1.png')
doc.add_picture(img_path_1, width=Inches(6.2))
caption('Figure 3-1: System Architecture — Offline-first 5-layer AI AFI Screening pipeline, '
        'from data input to local storage and optional HIS sync.')

doc.add_paragraph()
img_path_2 = os.path.join(FOLDER, 'word', 'media', 'image2.png')
doc.add_picture(img_path_2, width=Inches(5.2))
caption('Figure 3-2: Patient Journey — from CHW assessment to record and automated follow-up. '
        'Steps marked † require internet connectivity.')

h2('3.2  Research Methodology')
tbl4 = doc.add_table(rows=5, cols=4)
tbl4.style = 'Table Grid'
method_data = [
    ('Phase', 'Activities', 'Duration', 'Deliverables'),
    ('Phase 0\nData & Infrastructure',
     'Submit IRB (MFU + CCRU/OxTREC) | Analyse CCRU retrospective AFI data (de-identified) | '
     'Define disease prevalence baseline | Draft Data Sharing Agreement',
     'Months 1–3',
     'IRB approval + AFI dataset schema + DSA signed'),
    ('Phase 1\nModel Development',
     'Fine-tune Whisper for Thai medical domain | Build AFI Clinical Schema | '
     'Develop risk model on CCRU retrospective data | Prototype mobile UI (LINE OA + web)',
     'Months 4–8',
     'System prototype v1 | WER ≤10% | Slot-fill accuracy ≥0.85'),
    ('Phase 2\nField Validation',
     'Deploy in 3–5 CCRU-network SHPHs | Validate on 150+ real AFI cases | '
     'Measure sensitivity/specificity and usability | Test offline and connected modes | '
     'Iterate model',
     'Months 9–15',
     'Clinical validation report | All KPIs met'),
    ('Phase 3\nDissemination',
     'Analyse results | Submit peer-reviewed publication | '
     'Produce policy brief for MoPH and HSRI | Release de-identified open dataset | '
     'Discuss Phase 2 funding (Wellcome Trust / NRCT)',
     'Months 16–18',
     '≥1 publication + policy brief + open dataset'),
]
for i, row_data in enumerate(method_data):
    for j, val in enumerate(row_data):
        tbl4.rows[i].cells[j].text = val
        tbl4.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    if i == 0:
        for cell in tbl4.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '993C1D')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        shade_cell(tbl4.rows[i].cells[0], 'FAECE7')
set_col_widths(tbl4, [3.0, 7.5, 2.5, 3.5])
doc.add_paragraph()

h2('3.2.1  CCRU Dataset for ML Model Development')
body(
    'CCRU/MORU holds a unique data asset that is central to the AI development in this project:'
)
bullet(
    'Retrospective AFI patient records: more than 500 cases from rural and highland communities '
    'in Chiang Rai, with confirmed diagnoses (PCR/serology) for scrub typhus, dengue, '
    'leptospirosis, and bacteraemia — forming the labelled training and validation dataset.'
)
bullet(
    'Disease prevalence priors: seasonal and geographic distribution data used to calibrate '
    'Bayesian Priors in the risk-scoring model to reflect local epidemiological context.'
)
bullet(
    'Clinical features and lab correlates: symptoms, vital signs, basic blood results, and '
    'outcomes for feature engineering and training set construction.'
)
bullet(
    'AMR and antibiotic-use patterns: antimicrobial susceptibility testing (AST) results and '
    'prescribing patterns for development of the Antibiotic Stewardship Signal.'
)
bullet(
    'Field network: established relationships with SHPH staff and CHWs in the target sub-districts, '
    'reducing barriers to prospective data collection in the field validation phase.'
)
body(
    'All CCRU data will be de-identified to HIPAA-equivalent standards before transfer to the '
    'MFU team, governed by a formal Data Sharing Agreement (DSA) executed in Months 1–2.'
)

h2('3.3  Study Population and Sites')
body(
    'Target population: patients presenting with acute febrile illness (temperature ≥38°C for '
    'no more than 2 weeks) at CCRU-network SHPHs in districts with high AFI incidence, '
    'Chiang Rai province.'
)
body(
    'Inclusion criteria: age ≥15 years | presenting with AFI | able to provide informed consent.'
)
body(
    'Exclusion criteria: patients who have already received antibiotics for ≥48 hours | '
    'patients unable to communicate.'
)
body(
    'Sites: 3–5 SHPHs in the CCRU Chiang Rai network (names to be confirmed in consultation '
    'with Dr. Carlo Perrone). Field validation will encompass sites both with and without stable '
    'internet connectivity to validate system performance in Offline and Connected modes.'
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 4: Collaboration and Roles')

tbl5 = doc.add_table(rows=5, cols=3)
tbl5.style = 'Table Grid'
collab_data = [
    ('Organisation', 'Primary Role', 'Contributions'),
    ('Mae Fah Luang University (MFU)',
     'Lead grant recipient | AI/NLP system development | Thai-language UI design | '
     'HIS integration architecture | Grant administration',
     'AI/ML research team | MoU with MoPH | IRB process | Infrastructure'),
    ('CCRU / MORU, Chiang Rai',
     '[1] ML Training Dataset: deliver de-identified AFI database (>500 cases) with confirmed '
     'diagnoses, lab results, and outcomes for training/validation and Bayesian Prior calibration\n'
     '[2] Clinical Trial Support: prospective data source for Phase 2 field validation; '
     'clinical expert review of model outputs; coordination of 150+ case collection across '
     'SHPH network\n'
     '[3] AMR and Clinical Expertise: design of antibiotic stewardship signal; field research '
     'assistants and travel costs (via MFU–MORU subcontract)',
     'AFI retrospective dataset (>500, labelled) | Disease prevalence priors | '
     'AMR/AST data | SHPH field network | Clinical validation site | '
     'Expert ground-truth review | Oxford/Wellcome affiliation'),
    ('Partner SHPHs (≥3 sites)',
     'Field testing sites | Real end-users (CHWs/nurses) | Prospective data collection',
     'Real AFI patients | Local context | User feedback'),
    ('Chiang Rai Provincial Health Office (PHO)',
     'Policy stakeholder | Scale-up pathway',
     'Institutional legitimacy | Policy pathway'),
]
for i, row_data in enumerate(collab_data):
    for j, val in enumerate(row_data):
        tbl5.rows[i].cells[j].text = val
        tbl5.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)
    if i == 0:
        for cell in tbl5.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '5F5E5A')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif i == 2:
        shade_cell(tbl5.rows[i].cells[0], 'FFF8ED')
        shade_cell(tbl5.rows[i].cells[1], 'FFF8ED')
        shade_cell(tbl5.rows[i].cells[2], 'FFF8ED')
set_col_widths(tbl5, [3.5, 7.5, 5.5])

doc.add_paragraph()
body(
    'Note: All research funding flows through MFU as the lead institution in compliance with '
    'HSRI regulations. Budget allocated to CCRU/MORU for field research assistants and travel '
    'is administered via a formal subcontract or Research Collaboration Agreement between '
    'MFU and MORU.'
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 5: Work Plan and Timeline')

tbl6 = doc.add_table(rows=7, cols=3)
tbl6.style = 'Table Grid'
timeline_data = [
    ('Month', 'Key Activities', 'Milestone'),
    ('1–3',
     'Submit IRB (MFU + CCRU/OxTREC) | Negotiate CCRU Data Sharing Agreement | '
     'Analyse retrospective data | Draft AFI Clinical Schema',
     'IRB approved | DSA signed | Schema v1'),
    ('4–6',
     'Fine-tune Whisper Thai medical ASR | Develop NER/slot-filler | '
     'Build risk model (Bayesian + ML) on retrospective data',
     'WER ≤10% | Slot-fill accuracy ≥0.85'),
    ('7–8',
     'Build mobile UI (LINE OA primary + web fallback) | Internal testing | '
     'User Acceptance Testing (UAT) with CHWs/nurses | UX refinement',
     'System prototype v1 | Internal usability score ≥65'),
    ('9–12',
     'Deploy at 2 SHPHs (soft field trial) | Collect 75 cases | '
     'Measure sensitivity/specificity Round 1 | Go/No-go review',
     'Mid-field-trial report | Model iteration'),
    ('13–15',
     'Expand to ≥3 SHPHs | Complete 150+ cases | '
     'Measure AMR stewardship outcomes | Final clinical validation',
     'Clinical validation report | All KPIs met'),
    ('16–18',
     'Analyse and write up results | Submit peer-reviewed paper | '
     'Produce policy brief | Release open dataset | '
     'Discuss Phase 2 funding (Wellcome Trust / NRCT)',
     '≥1 publication | Policy brief | Public dataset'),
]
for i, row_data in enumerate(timeline_data):
    for j, val in enumerate(row_data):
        tbl6.rows[i].cells[j].text = val
        tbl6.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
    if i == 0:
        for cell in tbl6.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '0F6E56')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif i % 2 == 0:
        for cell in tbl6.rows[i].cells:
            shade_cell(cell, 'E8F5F0')
set_col_widths(tbl6, [1.8, 10.0, 4.7])
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 6: Indicative Budget')

tbl7 = doc.add_table(rows=10, cols=3)
tbl7.style = 'Table Grid'
budget_data = [
    ('Budget Category', 'Details', 'Amount (THB)'),
    ('Personnel',
     'PI, Co-PI, 2 AI/ML researchers, 2 research assistants (18 months)',
     '1,800,000'),
    ('Cloud & AI Infrastructure',
     'GPU compute (AWS/GCP), LLM API, data storage, deployment',
     '800,000'),
    ('Field Operations',
     'SHPH travel, field equipment, CHW and nurse training',
     '600,000'),
    ('Materials & Software',
     'Field trial devices, software licences, data annotation tools',
     '400,000'),
    ('Data Processing & Analysis',
     'Data annotation, statistical analysis, external expert review',
     '300,000'),
    ('Dissemination',
     'Open-access publication fees, conference, policy dissemination',
     '200,000'),
    ('CCRU/MORU Support\n(Field RA + Travel)',
     '1 field research assistant × 12 months (25,000/month) + field travel\n'
     'Via MFU–MORU subcontract',
     '450,000'),
    ('Institutional Overhead (10%)', '', '455,000'),
    ('TOTAL', '', '5,005,000'),
]
for i, row_data in enumerate(budget_data):
    for j, val in enumerate(row_data):
        tbl7.rows[i].cells[j].text = val
        tbl7.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
    if i == 0:
        for cell in tbl7.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '854F0B')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif i == 9:
        for cell in tbl7.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, 'FFF3CD')
    elif i % 2 == 0:
        for cell in tbl7.rows[i].cells:
            shade_cell(cell, 'FFF8ED')
set_col_widths(tbl7, [4.5, 9.0, 3.0])

doc.add_paragraph()
body(
    'Note: The above figures are indicative estimates. Final amounts will be aligned with the '
    'official HSRI budget template upon submission.'
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 7: Expected Outputs, Outcomes, and Impact')

h2('7.1  Outputs')
bullet('Open-source Thai-language AI AFI screening tool (mobile app: LINE OA + web)')
bullet('Fine-tuned Thai medical speech recognition model (publicly released)')
bullet('De-identified AFI patient dataset from Chiang Rai (open data release)')
bullet('At least one peer-reviewed journal article')
bullet('Policy brief for the Ministry of Public Health and HSRI')

h2('7.2  Outcomes')
bullet(
    'CHWs and SHPH nurses equipped with a practical, field-ready decision-support tool that '
    'works without internet access'
)
bullet('Reduced late diagnosis of scrub typhus and sepsis at pilot sites')
bullet('Reduced inappropriate antibiotic prescribing in AFI cases by at least 15%')
bullet(
    'Foundation established for scale-up through the MoPH Digital Health network in a '
    'subsequent phase'
)

h2('7.3  Long-term Impact')
bullet(
    'Reduced health inequity for vulnerable populations in remote northern Thailand'
)
bullet('Support for the National AMR Strategic Plan 2023–2027')
bullet(
    'Replicable model for Thai-language AI health tools for infectious diseases, '
    'scalable to the GMS/ASEAN region'
)
bullet(
    'Pathway to further funding: Wellcome Trust, NRCT (บพข.), or Gates Foundation'
)
bullet(
    'Regulatory pathway: Phase 2 will pursue FDA (Thailand) registration as a Software as a '
    'Medical Device (SaMD) to enable formal adoption in the public health system'
)
bullet(
    'NHSO (สปสช.) benefit package pathway: following clinical validation and SaMD registration, '
    'the system will be proposed for inclusion in the NHSO benefit package as a '
    'primary-level digital health service, securing long-term financial sustainability'
)

h2('7.4  Communication and Utilisation Plan')
body(
    'Audience 1 — Policy makers (MoPH, HSRI, Chiang Rai PHO): two policy briefs (Months 12 '
    'and 18); presentation to the Chiang Rai PHO and MoPH Bureau of Policy and Strategy for '
    'consideration of provincial and national scale-up.'
)
body(
    'Audience 2 — Health professionals and researchers: at least one peer-reviewed publication '
    '(Months 16–18); presentation at a national or international academic conference; '
    'open dataset release for the research community.'
)
body(
    'Audience 3 — Communities and CHWs: Training of Trainers (ToT) workshop for CHWs and '
    'SHPH nurses at field sites, enabling independent system operation after project completion.'
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 8: Research Ethics')
body(
    'This project involves human subjects research and will be conducted under full ethical '
    'oversight as follows:'
)
bullet(
    'Ethics approval from the MFU Human Research Ethics Committee will be obtained before '
    'any data collection commences.'
)
bullet(
    'Ethics review from the CCRU/MORU Ethics Committee (Oxford Tropical Research Ethics '
    'Committee: OxTREC, or equivalent) will be sought in parallel.'
)
bullet(
    'All research participants will provide written informed consent prior to enrolment.'
)
bullet(
    'All patient data will be de-identified before entering any AI processing pipeline; '
    'no personally identifiable information will be used.'
)
bullet(
    'A formal Data Sharing Agreement (DSA) between MFU and CCRU/MORU will be executed '
    'before the project commences.'
)
bullet(
    'The AI system is a decision-support tool, not an autonomous diagnostic system, '
    'which mitigates patient safety risk.'
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 9
# ═════════════════════════════════════════════════════════════════════════════
h1('Section 9: Research Team')

tbl8 = doc.add_table(rows=5, cols=4)
tbl8.style = 'Table Grid'
team_data = [
    ('Name', 'Affiliation', 'Role', 'Expertise'),
    ('Assoc. Prof. Dr. Nacha Choldarongkul',
     'MFU, School of ADT',
     'Principal Investigator',
     'AI/ML, Health Informatics, NLP, Software Engineering'),
    ('Asst. Prof. Dr. Phattaramon Wuttipitayamongkol',
     'MFU, School of ADT',
     'Co-Investigator (Co-PI)',
     'AI/ML, Digital Health, Health Informatics'),
    ('Dr. Carlo Perrone',
     'CCRU / MORU, Chiang Rai',
     'Co-Investigator',
     'Acute Febrile Illness, AMR, Scrub Typhus, Tropical Medicine'),
    ('[Clinical Collaborator TBC]',
     'Chiang Rai PHO / Chiang Rai Hospital',
     'Clinical Advisor',
     'Infectious Disease, Primary Care'),
]
for i, row_data in enumerate(team_data):
    for j, val in enumerate(row_data):
        tbl8.rows[i].cells[j].text = val
        tbl8.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10.5)
    if i == 0:
        for cell in tbl8.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True
            shade_cell(cell, '534AB7')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif i % 2 == 0:
        for cell in tbl8.rows[i].cells:
            shade_cell(cell, 'EEEDFE')
set_col_widths(tbl8, [4.5, 3.8, 3.5, 4.7])
doc.add_paragraph()

# ── Footer note ──────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Submit via NRIIS: www.nriis.go.th  |  Deadline: 30 April 2026\n'
    'HSRI enquiries: 02 027 9701 ext. 9056 (Ms. Piyachat Somthong)'
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT_FILE)
print('Saved:', OUT_FILE)
