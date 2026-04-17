# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'G:\My Drive\Research\MORU\Research_Proposal_EN.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

for i, table in enumerate(doc.tables):
    print(f"=== TABLE {i} ===")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {j}: {cells}")
    print()
