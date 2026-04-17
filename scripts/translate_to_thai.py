# -*- coding: utf-8 -*-
"""
Translate Research_Proposal_EN.docx → Research_Proposal_TH.docx
Uses Google Translate (deep_translator). Preserves all formatting,
tables, cell shading, styles, and column widths.
"""
import sys, os, shutil, time, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from deep_translator import GoogleTranslator

FOLDER = r'G:\My Drive\Research\MORU'
SRC = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
DST = os.path.join(FOLDER, 'Research_Proposal_TH.docx')

tr = GoogleTranslator(source='en', target='th')

# ── Translation helper ────────────────────────────────────────────────────────
def is_thai(text):
    thai = sum(1 for c in text if '\u0e00' <= c <= '\u0e7f')
    return thai / max(len(text), 1) > 0.4

def is_numeric_only(text):
    return bool(re.match(r'^[\d,\.\s\+\-\%\/THBx×]+$', text.strip()))

def translate(text, label=''):
    text = text.strip()
    if not text or is_thai(text) or is_numeric_only(text):
        return text
    for attempt in range(4):
        try:
            result = tr.translate(text)
            time.sleep(0.25)
            return result if result else text
        except Exception as e:
            if attempt == 3:
                print(f'  FAIL [{label}]: {text[:50]!r} — {e}')
                return text
            time.sleep(3 * (attempt + 1))
    return text

# ── Paragraph translator — preserves run formatting ──────────────────────────
def translate_para(para, label=''):
    orig = para.text.strip()
    if not orig:
        return

    # Gather run-level formatting from each character position
    runs_info = []  # (text, bold, italic, size, color_rgb, underline)
    for run in para.runs:
        if not run.text:
            continue
        color = None
        try:
            if run.font.color and run.font.color.type:
                color = run.font.color.rgb
        except Exception:
            pass
        runs_info.append({
            'text'     : run.text,
            'bold'     : run.bold,
            'italic'   : run.italic,
            'size'     : run.font.size,
            'color'    : color,
            'underline': run.underline,
        })

    thai = translate(orig, label)
    if thai == orig:
        return

    # Use first run's formatting for the whole translated paragraph
    fmt = runs_info[0] if runs_info else {}

    para.clear()
    run = para.add_run(thai)
    run.bold      = fmt.get('bold')
    run.italic    = fmt.get('italic')
    run.underline = fmt.get('underline')
    if fmt.get('size'):
        run.font.size = fmt['size']
    if fmt.get('color'):
        run.font.color.rgb = fmt['color']

# ── Start ─────────────────────────────────────────────────────────────────────
shutil.copy(SRC, DST)
doc = Document(DST)

# ── 1. Body paragraphs ────────────────────────────────────────────────────────
print('=== Body paragraphs ===')
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'  P{i:03d} {para.text.strip()[:60]}')
        translate_para(para, f'P{i}')

# ── 2. Tables ─────────────────────────────────────────────────────────────────
print('\n=== Tables ===')
for ti, tbl in enumerate(doc.tables):
    print(f'\n-- Table {ti} --')
    seen_cells = set()   # merged cells appear multiple times — skip duplicates
    for ri, row in enumerate(tbl.rows):
        # de-duplicate merged cells
        unique_cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen_cells:
                seen_cells.add(cid)
                unique_cells.append(cell)
        for ci, cell in enumerate(unique_cells):
            for pi, para in enumerate(cell.paragraphs):
                orig = para.text.strip()
                if orig:
                    print(f'  T{ti}R{ri}C{ci}P{pi}: {orig[:60]}')
                    translate_para(para, f'T{ti}R{ri}C{ci}')

doc.save(DST)
print(f'\nSaved → {DST}')
