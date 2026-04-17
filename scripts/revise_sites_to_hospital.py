# -*- coding: utf-8 -*-
"""
Revise Research_Proposal_EN.docx based on Carlo's feedback:
Move prospective validation data collection from SHPHs to district hospitals.
Rationale: SHPH staff don't draw blood (needed for confirmed diagnosis);
CCRU's confirmed-case dataset was collected at district hospitals.
The AI tool target (SHPHs) stays the same; only Phase 2 validation sites change.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)

def rewrite_para(para, new_text):
    """Replace paragraph text, preserving first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    f = para.runs[0]
    bold, italic, size, underline = f.bold, f.italic, f.font.size, f.underline
    color = None
    try:
        if f.font.color and f.font.color.type:
            color = f.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(new_text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

def patch_para_containing(para, find, replace):
    """Replace a substring within a paragraph's text, preserving formatting."""
    if find in para.text:
        new_text = para.text.replace(find, replace)
        print(f'  PATCH: {repr(find[:60])} → {repr(replace[:60])}')
        rewrite_para(para, new_text)
        return True
    return False

changes = 0

# ── Body paragraphs ────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    txt = para.text

    # 3.3 Target population: SHPH → district hospital
    if 'at CCRU-network SHPHs in districts with high AFI incidence' in txt:
        new = txt.replace(
            'at CCRU-network SHPHs in districts with high AFI incidence',
            'at district hospitals within the CCRU Chiang Rai network'
        )
        rewrite_para(para, new)
        print(f'P(population): updated to district hospitals')
        changes += 1

    # 3.3 Sites line: change SHPHs to district hospitals + add rationale
    if txt.strip().startswith('Sites: 3') and 'SHPH' in txt:
        new = (
            'Sites: 3–5 district hospitals in the CCRU Chiang Rai network '
            '(names to be confirmed in consultation with Dr. Carlo Perrone). '
            'District hospitals are selected because staff routinely perform '
            'blood draws enabling gold-standard confirmatory diagnosis '
            '(PCR/serology/culture), consistent with how the CCRU retrospective '
            'dataset was collected. Following successful validation, the tool '
            'will be adapted and deployed at SHPH level.'
        )
        rewrite_para(para, new)
        print(f'P(sites): updated to district hospitals + rationale')
        changes += 1

    # Remove/update the internet connectivity note (now less relevant for hospitals)
    if 'Field validation will encompass sites both with and without stable internet' in txt:
        new = txt.replace(
            'Field validation will encompass sites both with and without stable internet connectivity',
            'Field validation sites will include hospitals both with and without stable internet connectivity, ensuring the offline-first architecture is tested under realistic conditions'
        )
        rewrite_para(para, new)
        print(f'P(connectivity): updated')
        changes += 1

    # Section 3.2 field network note — update prospective data collection reference
    if 'Field network: established relationships with SHPH staff' in txt:
        new = txt.replace(
            'Field network: established relationships with SHPH staff and CHWs in the target sub-districts, reducing barriers to prospective data collection in the field validation phase',
            'Field network: established relationships with district hospital staff and CHWs in the CCRU Chiang Rai network, enabling prospective data collection with confirmatory diagnostics in the field validation phase'
        )
        rewrite_para(para, new)
        print(f'P(field network): updated')
        changes += 1

# ── Tables ─────────────────────────────────────────────────────────────────────
seen_cells = set()
for ti, table in enumerate(doc.tables):
    seen_cells = set()
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen_cells:
                continue
            seen_cells.add(cid)

            for para in cell.paragraphs:
                txt = para.text

                # Table 3 Phase 2: SHPHs → district hospitals
                if 'Deploy in 3' in txt and 'CCRU-network SHPHs' in txt:
                    new = txt.replace(
                        'Deploy in 3–5 CCRU-network SHPHs',
                        'Deploy at 3–5 CCRU-network district hospitals'
                    ).replace(
                        'Validate on 150+ real AFI cases',
                        'Validate on 150+ real AFI cases with blood-confirmed diagnoses'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: Phase 2 activities updated')
                    changes += 1

                # Table 3 Phase 2 deliverables — no change needed

                # Table 4 CCRU contributions — update SHPH network reference
                if 'coordination of 150+ case collection across SHPH network' in txt:
                    new = txt.replace(
                        'coordination of 150+ case collection across SHPH network',
                        'coordination of 150+ case collection across district hospital network (blood draw / confirmed diagnosis)'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: CCRU contributions updated')
                    changes += 1

                if 'prospective data source for Phase 2 field validation' in txt and 'SHPH' in txt:
                    new = txt.replace(
                        'Clinical Trial Support: prospective data source for Phase 2 field validation; clinical expert review of model outputs; coordination of 150+ case collection across SHPH network',
                        'Clinical Trial Support: prospective data source for Phase 2 field validation at district hospitals (blood draw for confirmed diagnosis); clinical expert review of model outputs; coordination of 150+ case collection'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: CCRU clinical trial support updated')
                    changes += 1

                # Table 4 Partner SHPHs row — clarify their role as future deployment, not validation
                if 'Partner SHPHs' in txt or ('Field testing sites' in txt and 'Prospective data collection' in txt):
                    if 'Prospective data collection' in txt:
                        new = txt.replace(
                            'Prospective data collection',
                            'Future deployment sites post-validation'
                        )
                        rewrite_para(para, new)
                        print(f'T{ti}R{ri}: Partner SHPHs role clarified')
                        changes += 1

                # Table 5 timeline — months 9-12
                if 'Deploy at 2 SHPHs' in txt:
                    new = txt.replace(
                        'Deploy at 2 SHPHs (soft field trial)',
                        'Deploy at 2 district hospitals (soft field trial)'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: Timeline months 9-12 updated')
                    changes += 1

                # Table 5 timeline — months 13-15
                if 'Expand to' in txt and '≥3 SHPHs' in txt:
                    new = txt.replace(
                        'Expand to ≥3 SHPHs',
                        'Expand to ≥3 district hospitals'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: Timeline months 13-15 updated')
                    changes += 1

                # KPI table (Table 2) — update "From ≥3 SHPH partner sites"
                if 'From ≥3 SHPH partner sites' in txt:
                    new = txt.replace(
                        'From ≥3 SHPH partner sites',
                        'From ≥3 district hospital partner sites'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: KPI measurement method updated')
                    changes += 1

                # Cover table (Table 0) — Study Sites row
                if 'CCRU/MORU + 3' in txt and 'SHPH' in txt:
                    new = txt.replace(
                        'Chiang Rai: CCRU/MORU + 3–5 partner Sub-district Health Promotion Hospitals (SHPHs)',
                        'Chiang Rai: CCRU/MORU + 3–5 partner district hospitals (prospective validation); target SHPH deployment post-validation'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: Cover table study sites updated')
                    changes += 1

                # Table 1 alignment — "1 AI AFI screening tool deployed in ≥3 SHPHs"
                if '1 AI AFI screening tool deployed in' in txt and 'SHPHs' in txt:
                    new = txt.replace(
                        '≥3 SHPHs',
                        '≥3 district hospitals (validation); SHPH deployment in Phase 2'
                    )
                    rewrite_para(para, new)
                    print(f'T{ti}R{ri}: HSRI alignment table updated')
                    changes += 1

print(f'\nTotal changes: {changes}')
doc.save(TMP)
os.replace(TMP, DOCX)
print('Done — saved.')
