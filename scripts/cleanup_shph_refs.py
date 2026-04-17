# -*- coding: utf-8 -*-
"""Final cleanup: fix remaining SHPH references that need updating."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)

def rewrite_para(para, new_text):
    if not para.runs:
        para.add_run(new_text); return
    f = para.runs[0]
    bold, italic, size, underline = f.bold, f.italic, f.font.size, f.underline
    color = None
    try:
        if f.font.color and f.font.color.type:
            color = f.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(new_text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

changes = 0

# ── Body paragraphs ────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    txt = para.text.strip()

    # P058: Sites paragraph — partially broken by collision, rewrite cleanly
    if txt.startswith('Sites: 3') and 'SHPH' in txt:
        new = (
            'Sites: 3–5 district hospitals in the CCRU Chiang Rai network '
            '(names to be confirmed in consultation with Dr. Carlo Perrone). '
            'District hospitals are selected because staff routinely perform '
            'blood draws enabling gold-standard confirmatory diagnosis '
            '(PCR/serology/culture), consistent with how the CCRU retrospective '
            'dataset was collected. Field validation sites include hospitals both '
            'with and without stable internet connectivity, ensuring the '
            'offline-first architecture is tested under realistic conditions. '
            'Following successful validation, the tool will be adapted and '
            'deployed at SHPH level.'
        )
        rewrite_para(para, new)
        print(f'P(sites): fixed → district hospitals + rationale')
        changes += 1

    # P024: Objective 2.4 — update evaluation site
    if 'usability and clinical accuracy' in txt and 'CCRU-network SHPHs' in txt:
        new = txt.replace(
            'at CCRU-network SHPHs',
            'at CCRU-network district hospitals (Phase 2 field validation)'
        )
        rewrite_para(para, new)
        print(f'P024: objective updated to district hospitals')
        changes += 1

# ── Cover table (Table 0, R7C1P0) ─────────────────────────────────────────────
cell = doc.tables[0].rows[7].cells[1]
for para in cell.paragraphs:
    if 'SHPH' in para.text and 'CCRU/MORU' in para.text:
        new = (
            'Chiang Rai: CCRU/MORU + 3–5 partner district hospitals '
            '(Phase 2 prospective validation) + MFU; '
            'SHPH deployment planned post-validation'
        )
        rewrite_para(para, new)
        print(f'T0R7C1: cover table study sites fixed')
        changes += 1

# ── Table 4 R3C0P0: Partner SHPHs → Partner District Hospitals ────────────────
cell = doc.tables[4].rows[3].cells[0]
for para in cell.paragraphs:
    if 'Partner SHPHs' in para.text:
        new = para.text.replace('Partner SHPHs (≥3 sites)', 'Partner District Hospitals (≥3 sites)')
        rewrite_para(para, new)
        print(f'T4R3C0: Partner SHPHs → Partner District Hospitals')
        changes += 1

print(f'\nTotal changes: {changes}')
doc.save(TMP)
os.replace(TMP, DOCX)
print('Done — saved.')
