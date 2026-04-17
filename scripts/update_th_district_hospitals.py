# -*- coding: utf-8 -*-
"""
Update Research_Proposal_TH.docx to reflect Carlo's revision:
prospective validation moved from SHPHs → district hospitals (โรงพยาบาลระดับอำเภอ).
All translations hand-crafted to match Thai health-system terminology.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_TH.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_TH_tmp.docx')

doc = Document(DOCX)

def rewrite_para(para, new_text):
    """Replace paragraph text preserving first run's formatting."""
    if not para.runs:
        para.add_run(new_text); return
    f = para.runs[0]
    bold, italic, size, underline = f.bold, f.italic, f.font.size, f.underline
    color = None
    try:
        if f.font.color and f.font.color.type:
            color = f.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(new_text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

def fix_body(idx, new_text):
    para = doc.paragraphs[idx]
    print(f'  Body P{idx:03d}: {para.text[:60]!r}')
    rewrite_para(para, new_text)

def fix_table(ti, ri, ci, pi, new_text):
    cell = doc.tables[ti].rows[ri].cells[ci]
    para = cell.paragraphs[pi]
    print(f'  T{ti}R{ri}C{ci}P{pi}: {para.text[:60]!r}')
    rewrite_para(para, new_text)

print('=== Updating Thai document ===\n')

# ── Body paragraphs ────────────────────────────────────────────────────────────

# P024 — Objective 2.4: evaluation site → district hospitals
fix_body(24,
    'ประเมินการใช้งานและความแม่นยำทางคลินิกของระบบภายใต้สภาวะการใช้งานจริง'
    'ที่โรงพยาบาลระดับอำเภอในเครือข่าย CCRU (การตรวจสอบภาคสนาม ระยะที่ 2)')

# P052 — Field network: SHPH staff → district hospital staff
fix_body(52,
    'เครือข่ายภาคสนาม: สร้างความสัมพันธ์กับบุคลากรโรงพยาบาลระดับอำเภอและ อสม. '
    'ในเครือข่าย CCRU เชียงราย เพื่อรองรับการเก็บข้อมูลเชิงคาดการณ์ด้วยการวินิจฉัยยืนยัน '
    'ในระยะตรวจสอบภาคสนาม')

# P055 — Target population: at district hospitals
fix_body(55,
    'ประชากรเป้าหมาย: ผู้ป่วยที่มีอาการไข้เฉียบพลัน (อุณหภูมิ ≥38°C เป็นเวลาไม่เกิน 2 สัปดาห์) '
    'ที่เข้ารับบริการที่โรงพยาบาลระดับอำเภอในเครือข่าย CCRU จังหวัดเชียงราย')

# P058 — Sites: 3–5 district hospitals + full rationale
fix_body(58,
    'สถานที่: โรงพยาบาลระดับอำเภอ 3–5 แห่งในเครือข่าย CCRU เชียงราย '
    '(ต้องยืนยันชื่อในการหารือกับ Dr. Carlo Perrone) '
    'โรงพยาบาลระดับอำเภอได้รับเลือกเนื่องจากบุคลากรสามารถเจาะเลือดได้เป็นประจำ '
    'ทำให้สามารถวินิจฉัยยืนยันด้วยมาตรฐานอ้างอิง (PCR/ซีรัมวิทยา/การเพาะเชื้อ) '
    'ซึ่งสอดคล้องกับวิธีการเก็บข้อมูลย้อนหลังของ CCRU '
    'สถานที่ตรวจสอบครอบคลุมโรงพยาบาลทั้งที่มีและไม่มีการเชื่อมต่ออินเทอร์เน็ตที่เสถียร '
    'เพื่อทดสอบสถาปัตยกรรม Offline-first ในสภาพแวดล้อมจริง '
    'หลังการตรวจสอบสำเร็จ เครื่องมือจะปรับและขยายการใช้งานไปยังระดับ รพ.สต.')

print()

# ── Tables ─────────────────────────────────────────────────────────────────────

# Table 0 R7C1 — Cover table: study sites
fix_table(0, 7, 1, 0,
    'เชียงราย: CCRU/MORU + โรงพยาบาลระดับอำเภอพันธมิตร 3–5 แห่ง '
    '(การตรวจสอบเชิงคาดการณ์ ระยะที่ 2) + MFU; '
    'วางแผนขยายไปยัง รพ.สต. หลังการตรวจสอบ')

# Table 1 R5C1 — HSRI alignment: deployed in district hospitals
fix_table(1, 5, 1, 0,
    'เครื่องมือคัดกรอง AI AFI จำนวน 1 ชุด ใช้งานใน ≥3 โรงพยาบาลระดับอำเภอ (การตรวจสอบ); '
    'ขยายไปยัง รพ.สต. ในระยะที่ 2; ≥150 ราย; '
    'ประชากรในพื้นที่ดูแล ≥10,000 ราย ในระยะขยายผล')

# Table 2 R6C2 — KPI: from district hospital sites
fix_table(2, 6, 2, 0,
    'จากสถานพยาบาลระดับอำเภอพันธมิตร ≥3 แห่ง')

# Table 3 R3C1 — Work plan Phase 2: district hospitals + blood-confirmed
fix_table(3, 3, 1, 0,
    'ปรับใช้ที่โรงพยาบาลระดับอำเภอในเครือข่าย CCRU 3–5 แห่ง | '
    'ตรวจสอบกรณี AFI จริงมากกว่า 150 กรณี พร้อมการวินิจฉัยยืนยันด้วยเลือด (PCR/ซีรัมวิทยา) | '
    'วัดความไว/ความจำเพาะและการใช้งาน | '
    'ทดสอบโหมดออฟไลน์และการเชื่อมต่อ | '
    'ปรับปรุงโมเดล')

# Table 4 R2C1 — CCRU primary role: district hospitals for field validation
fix_table(4, 2, 1, 0,
    '[1] ชุดข้อมูลการฝึกอบรม ML: ส่งมอบฐานข้อมูล AFI ที่ไม่ระบุตัวตน (>500 ราย) '
    'พร้อมการวินิจฉัยยืนยัน ผลตรวจทางห้องปฏิบัติการ และผลลัพธ์ '
    'สำหรับการฝึกและตรวจสอบโมเดล และการสอบเทียบ Bayesian Prior\n'
    '[2] การสนับสนุนการทดลองทางคลินิก: แหล่งข้อมูลเชิงคาดการณ์ '
    'สำหรับการตรวจสอบภาคสนาม ระยะที่ 2 ที่โรงพยาบาลระดับอำเภอ '
    '(ซึ่งการเจาะเลือดช่วยยืนยันการวินิจฉัย); '
    'ผู้เชี่ยวชาญทางคลินิกตรวจทานผลโมเดล; ประสานการเก็บข้อมูลมากกว่า 150 ราย\n'
    '[3] ความเชี่ยวชาญด้าน AMR และคลินิก: ออกแบบสัญญาณการดูแลการใช้ยาปฏิชีวนะ; '
    'ผู้ช่วยวิจัยภาคสนามและค่าเดินทาง (ผ่านสัญญาช่วง MFU–MORU)')

# Table 4 R2C2 — CCRU contributions: district hospital field network
fix_table(4, 2, 2, 0,
    'ชุดข้อมูลย้อนหลัง AFI (>500 มีป้ายกำกับ) | ข้อมูลความชุกของโรค | '
    'ข้อมูล AMR/AST | เครือข่ายโรงพยาบาลระดับอำเภอ | '
    'สถานที่ตรวจสอบทางคลินิก | การตรวจสอบโดยผู้เชี่ยวชาญ | '
    'สังกัดออกซ์ฟอร์ด/เวลคัม')

# Table 4 R3C0 — Partner org name: district hospitals
fix_table(4, 3, 0, 0,
    'โรงพยาบาลระดับอำเภอพันธมิตร (≥3 แห่ง)')

# Table 4 R3C1 — Partner role: future deployment sites
fix_table(4, 3, 1, 0,
    'สถานที่ขยายผลในอนาคต (หลังการตรวจสอบ) | '
    'ผู้ใช้งานจริง (อสม./พยาบาล) | '
    'การมีส่วนร่วมของชุมชนเพื่อการขยายผล')

# Table 5 R4C1 — Timeline months 9–12: district hospitals
fix_table(5, 4, 1, 0,
    'ปรับใช้ที่โรงพยาบาลระดับอำเภอ 2 แห่ง (การทดลองภาคสนามเบื้องต้น) | '
    'เก็บข้อมูล 75 ราย พร้อมการวินิจฉัยยืนยันด้วยเลือด | '
    'วัดความไว/ความจำเพาะ รอบที่ 1 | '
    'ทบทวนผล Go/No-go')

# Table 5 R5C1 — Timeline months 13–15: district hospitals
fix_table(5, 5, 1, 0,
    'ขยายเป็น ≥3 โรงพยาบาลระดับอำเภอ | '
    'ครบ 150+ ราย | '
    'วัดผลลัพธ์การดูแลการใช้ยาปฏิชีวนะ AMR | '
    'การตรวจสอบทางคลินิกขั้นสุดท้าย')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nDone — saved.')
