# -*- coding: utf-8 -*-
"""
Add a List of Abbreviations section and expand first occurrences in body text.
"""
import sys, os, re, copy
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

# ── Master abbreviation dictionary ───────────────────────────────────────────
# Only terms that need definition (skip obvious ones like TOTAL, FY, TBC, LINE)
ABBREVS = {
    'ADT':   'Applied Digital Technology',
    'AFI':   'Acute Febrile Illness',
    'AI':    'Artificial Intelligence',
    'AMR':   'Antimicrobial Resistance',
    'API':   'Application Programming Interface',
    'ASEAN': 'Association of Southeast Asian Nations',
    'ASR':   'Automatic Speech Recognition',
    'AST':   'Antimicrobial Susceptibility Testing',
    'AWS':   'Amazon Web Services',
    'CCRU':  'Chiang Rai Clinical Research Unit',
    'CDSS':  'Clinical Decision Support System',
    'CHW':   'Community Health Worker',
    'DSA':   'Data Sharing Agreement',
    'EMR':   'Electronic Medical Record',
    'FDA':   'Food and Drug Administration (Thailand)',
    'FHIR':  'Fast Healthcare Interoperability Resources',
    'GCP':   'Google Cloud Platform',
    'GMS':   'Greater Mekong Subregion',
    'GPU':   'Graphics Processing Unit',
    'HIPAA': 'Health Insurance Portability and Accountability Act',
    'HIS':   'Hospital Information System',
    'HSRI':  'Health Systems Research Institute',
    'IRB':   'Institutional Review Board',
    'KPI':   'Key Performance Indicator',
    'LLM':   'Large Language Model',
    'MFU':   'Mae Fah Luang University',
    'ML':    'Machine Learning',
    'MoPH':  'Ministry of Public Health',
    'MORU':  'Mahidol Oxford Tropical Medicine Research Unit',
    'NER':   'Named Entity Recognition',
    'NHSO':  'National Health Security Office',
    'NLP':   'Natural Language Processing',
    'NRCT':  'National Research Council of Thailand',
    'NRIIS': 'National Research Information and Intelligence System',
    'PCR':   'Polymerase Chain Reaction',
    'PHO':   'Provincial Health Office',
    'PI':    'Principal Investigator',
    'RA':    'Research Assistant',
    'SHPH':  'Sub-district Health Promotion Hospital',
    'THB':   'Thai Baht',
    'TRL':   'Technology Readiness Level',
    'UAT':   'User Acceptance Testing',
    'UI':    'User Interface',
    'UX':    'User Experience',
    'WER':   'Word Error Rate',
}

# Abbreviations already defined inline in the cover table — skip expanding in body
SKIP_EXPAND = {'CCRU', 'MORU', 'MFU', 'HSRI', 'SHPH', 'CHW', 'AFI', 'AMR',
               'TRL', 'THB', 'PI', 'HIS', 'EMR', 'IRB', 'NER', 'WER',
               'CDSS', 'PHO'}

doc = Document(DOCX)

# ── 1. Insert "List of Abbreviations" after the cover table ─────────────────
# Find where Section 1 starts (first paragraph with "Section 1")
from docx.oxml import OxmlElement

body = doc.element.body
sec1_elem = None
for child in body:
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        from docx.text.paragraph import Paragraph
        p = Paragraph(child, doc)
        if p.text.strip().startswith('Section 1:'):
            sec1_elem = child
            break

if sec1_elem is None:
    print('ERROR: Section 1 not found'); sys.exit(1)

def make_para(text, bold=False, size=11, color=None, align='left', space_before=0, space_after=4):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    # spacing
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(space_before * 20)))
    sp.set(qn('w:after'),  str(int(space_after  * 20)))
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz');  sz.set(qn('w:val'), str(size * 2));  rPr.append(sz)
    sz2= OxmlElement('w:szCs');sz2.set(qn('w:val'), str(size * 2)); rPr.append(sz2)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    if color:
        cl = OxmlElement('w:color')
        cl.set(qn('w:val'), color)
        rPr.append(cl)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

# Section heading
sec1_elem.addprevious(make_para(''))   # blank line
sec1_elem.addprevious(make_para(''))   # blank line

# Build the abbreviations table XML
from docx.table import Table as DTable

def make_abbrev_table(doc):
    """Create a 2-column abbreviations table and return its XML element."""
    sorted_abbrevs = sorted(ABBREVS.items())
    rows = len(sorted_abbrevs) + 1  # +1 header
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.style = 'Table Grid'
    # Header
    hdr = tbl.rows[0]
    hdr.cells[0].text = 'Abbreviation'
    hdr.cells[1].text = 'Full Form'
    for cell in hdr.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement as OE
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OE('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0F6E56')
        tcPr.append(shd)
    # Rows
    for i, (abbr, full) in enumerate(sorted_abbrevs):
        row = tbl.rows[i + 1]
        row.cells[0].text = abbr
        row.cells[1].text = full
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10.5)
        if i % 2 == 0:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OE('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8F5F0')
                tcPr.append(shd)
        row.cells[0].paragraphs[0].runs[0].bold = True
    # Set column widths
    for row in tbl.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(13.0)
    return tbl._tbl

# Add heading and table into a temp doc, then move into position
heading_p = make_para('List of Abbreviations', bold=True, size=13,
                      color='0F6E56', space_before=6, space_after=4)
sec1_elem.addprevious(heading_p)

# We need to insert the table element before sec1_elem
# Temporarily add at end of doc, then move
tbl_elem = make_abbrev_table(doc)
# The table was just appended at end of body; we need to move it before sec1_elem
body.remove(tbl_elem)
sec1_elem.addprevious(tbl_elem)
print('Inserted List of Abbreviations table.')

# ── 2. Expand abbreviations on first use in body paragraphs ─────────────────
# Track which have been expanded
expanded = set(SKIP_EXPAND)   # these are already defined in cover/title area

def expand_text(text, expanded):
    """Replace first occurrence of each unexpanded abbreviation with Full (ABBR)."""
    changed = False
    for abbr, full in ABBREVS.items():
        if abbr in expanded:
            continue
        # Match as a whole word, not already in "Full (ABBR)" pattern
        pattern = r'(?<!\()(?<!\w)' + re.escape(abbr) + r'(?!\w)(?!\s*=)(?![^(]*\))'
        if re.search(pattern, text):
            replacement = f'{full} ({abbr})'
            text = re.sub(pattern, replacement, text, count=1)
            expanded.add(abbr)
            changed = True
    return text, changed

# Walk body elements in order
skip_next = False
for child in list(body):
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        from docx.text.paragraph import Paragraph
        p = Paragraph(child, doc)
        # Skip title/header paragraphs and section headings
        txt = p.text
        if not txt.strip():
            continue
        if txt.strip().startswith('Section ') or txt.strip().startswith('List of Abbrev'):
            continue
        if txt.strip().startswith('Figure') or txt.strip().startswith('Note:'):
            continue

        new_txt, changed = expand_text(txt, expanded)
        if changed:
            # Replace all run text with new expanded text
            # Simplest: clear runs and set first run text
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_txt
            else:
                p.add_run(new_txt)

    elif tag == 'tbl':
        from docx.table import Table
        t = Table(child, doc)
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text
                    if not txt.strip():
                        continue
                    new_txt, changed = expand_text(txt, expanded)
                    if changed:
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = new_txt
                        else:
                            para.add_run(new_txt)

print(f'Expanded {len(expanded) - len(SKIP_EXPAND)} abbreviations inline.')

# ── 3. Save ──────────────────────────────────────────────────────────────────
doc.save(TMP)
os.replace(TMP, DOCX)
print('Saved:', DOCX)
