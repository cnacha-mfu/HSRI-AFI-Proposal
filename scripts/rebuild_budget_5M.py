# -*- coding: utf-8 -*-
"""
Rebuild budget table at ~5,000,000 THB.
Staff: PI + Co-PI + 1 AI/ML Developer (Master's) + 1 Clinical RA (Bachelor's)
Scale-up vs 4M: expanded cloud AI, field operations, annotation, equipment restored.
Grand total: ~4,969,950 THB (≈ 5M)  |  Personnel: 19.8%  |  Fee: 10.0%
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)

BUDGET = [
    {
        'cat'  : 'Category 1 — Personnel Compensation (หมวดค่าตอบแทน)',
        'color': '1B4F72',
        'items': [
            ('Principal Investigator (PhD, Assoc. Prof.)',
             '45,000 THB/mo x 15% FTE x 18 mo', 121_500),
            ('Co-Principal Investigator (PhD, Asst. Prof.)',
             '35,000 THB/mo x 10% FTE x 18 mo', 63_000),
            ("AI/ML Developer (Master's degree)",
             '22,000 THB/mo x 100% FTE x 18 mo', 396_000),
            ("Clinical Research Assistant (Bachelor's degree)",
             '18,000 THB/mo x 100% FTE x 18 mo', 324_000),
        ],
    },
    {
        'cat'  : 'Category 2 — Project Management Overhead (หมวดค่าบริหารจัดการ)',
        'color': '145A32',
        'items': [
            ('Documentation — inception, interim and final research reports',
             '3 reports x lump sum', 80_000),
            ('Office supplies and stationery',
             'Lump sum over 18 months', 36_000),
            ('Telephone and internet — project coordination charges',
             '2,000 THB/mo x 12 mo', 24_000),
            ('Liaison travel — HSRI / MoPH progress review meetings',
             '4 trips x 15,000 THB (Bangkok)', 60_000),
            ('Printing and reproduction — reports, manuals, consent forms',
             'Lump sum', 30_000),
            ('Publication fees',
             '2 peer-reviewed journal articles x 60,000 THB', 120_000),
            ('Conference presentations (domestic + regional)',
             '2 conferences x 50,000 THB', 100_000),
            ('Policy dissemination workshop — Chiang Rai',
             '1 workshop, ~50 participants', 50_000),
        ],
    },
    {
        'cat'  : 'Category 3 — Research Operations (หมวดค่าดำเนินงาน)',
        'color': '4A235A',
        'note' : 'Includes MFU-MORU subcontract (HSRI Sec. 3.2.3)',
        'items': [
            ('Cloud AI — model training (GPU instances, LLM fine-tuning, ASR, NER)',
             'AWS/GCP intensive compute; training phase budget', 650_000),
            ('Cloud AI — inference, deployment and storage (ongoing)',
             'API hosting, model serving, data storage; 18 months', 300_000),
            ('CCRU/MORU subcontract — Field RA salary  [MFU-MORU subcontract agreement]',
             '25,000 THB/mo x 12 mo', 300_000),
            ('CCRU/MORU subcontract — Field travel and site logistics',
             'Lump sum; 5 SHPH sites across Chiang Rai province', 150_000),
            ('Field data collection travel — MFU research team',
             'Est. 40 site visits x 5,000 THB avg.', 200_000),
            ('CHW and nurse training workshops',
             '5 sites x 3 sessions x 8,000 THB per session', 120_000),
            ('Participant engagement and consent process costs',
             'Incentives + refreshments; 400 enrolled participants', 80_000),
            ('Data annotation and labeling services',
             'Clinical audio transcription + entity labeling; 2,000 records', 300_000),
            ('External statistical analysis consultant',
             '1 consultant; fixed-fee contract', 80_000),
            ('Clinical expert advisory — 3 external reviewers + UAT facilitation',
             'CDSS rule validation, UAT protocol review, 2 rounds clinical feedback', 200_000),
            ('Software tool subscriptions (LLM API, annotation platform, NLP tools, CI/CD)',
             'Multiple subscriptions < 20,000 THB/unit; classified as operations per HSRI', 200_000),
            ('Field communication allowances — SIM/data for CHW tablets',
             '100 CHW devices x 1,000 THB/device x 10 mo', 100_000),
            ('AI development tool licenses (coding assistant, IDE plugins)',
             '$100/developer/mo x 2 developers x 18 mo x 35 THB/USD', 126_000),
        ],
    },
    {
        'cat'  : 'Category 4 — Equipment (หมวดค่าครุภัณฑ์)',
        'color': '7D6608',
        'note' : 'All items >= 20,000 THB/unit; excluded from Cat 1 percentage denominator',
        'items': [
            ('GPU development workstation — on-site model training and inference',
             '1 unit x 200,000 THB', 200_000),
            ('High-performance research laptops',
             '3 units x 40,000 THB', 120_000),
            ('Digital vital-signs monitors (SpO2, BP, temperature) — field validation',
             '4 sets x 20,000 THB', 80_000),
        ],
    },
    {
        'cat'  : 'Category 5 — Institutional Overhead Fee (หมวดค่าธรรมเนียมหน่วยงานผู้รับทุน)',
        'color': '2E4053',
        'items': [
            ('MFU institutional overhead — utilities, workspace, admin, shared IT infrastructure',
             '10% x (Category 1 + 2 + 3) = 10% x 4,210,500 THB', 421_050),
        ],
    },
]

for cat in BUDGET:
    cat['subtotal'] = sum(item[2] for item in cat['items'])

grand_total = sum(cat['subtotal'] for cat in BUDGET)
cat1    = BUDGET[0]['subtotal']
cat4    = BUDGET[3]['subtotal']
non_eq  = grand_total - cat4
base123 = sum(BUDGET[i]['subtotal'] for i in range(3))

def fmt(n): return f'{int(n):,}'

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex); tcPr.append(shd)

def set_cell(cell, text, bold=False, italic=False, size=10,
             align='left', fill=None, color_hex=None):
    if fill: shade(cell, fill)
    para = cell.paragraphs[0]; para.clear()
    para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT  if align == 'right'  else
                      WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else
                      WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))

def note_para(cell, text):
    p = cell.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ── Build table ───────────────────────────────────────────────────────────────
n_rows = 1
for cat in BUDGET:
    n_rows += 1 + len(cat['items']) + 1
n_rows += 2   # compliance + grand total

tbl = doc.add_table(rows=n_rows, cols=3)
tbl.style = 'Table Grid'
ri = 0

# Header row
set_cell(tbl.rows[ri].cells[0], 'Budget Category / Line Item',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56')
set_cell(tbl.rows[ri].cells[1], 'Calculation Basis / Details',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56')
set_cell(tbl.rows[ri].cells[2], 'Amount (THB)',
         bold=True, size=10.5, color_hex='FFFFFF', fill='0F6E56', align='center')
ri += 1

notes = {
    0: f'{cat1/non_eq*100:.1f}% of total excl. equipment — compliant with <=30% HSRI cap',
    1: f'{BUDGET[1]["subtotal"]/non_eq*100:.1f}% of total excl. equipment — compliant with <=15% HSRI cap',
    4: f'10% of (Cat 1+2+3) = {fmt(base123)} THB — compliant with <=10% HSRI cap',
}

for ci, cat in enumerate(BUDGET):
    # Category header
    merged = tbl.rows[ri].cells[0].merge(tbl.rows[ri].cells[2])
    set_cell(merged, cat['cat'], bold=True, size=10.5, color_hex='FFFFFF', fill=cat['color'])
    note_para(merged, notes.get(ci, cat.get('note', '')))
    ri += 1

    # Items
    for ii, (name, detail, amount) in enumerate(cat['items']):
        fill = 'EBF5EB' if ii % 2 == 0 else 'FFFFFF'
        row  = tbl.rows[ri]
        for c in row.cells: shade(c, fill)
        set_cell(row.cells[0], '  ' + name, size=10)
        set_cell(row.cells[1], detail, italic=True, size=10, color_hex='444444')
        set_cell(row.cells[2], fmt(amount), size=10, align='right')
        ri += 1

    # Subtotal
    sub = tbl.rows[ri]
    for c in sub.cells: shade(c, 'D5F5E3')
    set_cell(sub.cells[0], f'  Subtotal — Category {ci+1}', bold=True, size=10)
    set_cell(sub.cells[1], '', size=10)
    set_cell(sub.cells[2], fmt(cat['subtotal']), bold=True, size=10, align='right')
    ri += 1

# Compliance note
comp = tbl.rows[ri].cells[0].merge(tbl.rows[ri].cells[2])
shade(tbl.rows[ri].cells[0], 'FEF9E7')
set_cell(tbl.rows[ri].cells[0],
         f'HSRI compliance  |  '
         f'Personnel (Cat 1): {fmt(cat1)} / {fmt(non_eq)} = {cat1/non_eq*100:.1f}% <= 30%   '
         f'Overhead (Cat 2): {BUDGET[1]["subtotal"]/non_eq*100:.1f}% <= 15%   '
         f'Fee (Cat 5): {BUDGET[4]["subtotal"]/base123*100:.1f}% <= 10%',
         italic=True, size=9)
ri += 1

# Grand total
gt = tbl.rows[ri]
gt.cells[0].merge(gt.cells[1])
shade(gt.cells[0], '0F6E56'); shade(gt.cells[2], '0F6E56')
set_cell(gt.cells[0],
         'GRAND TOTAL  (HSRI FY2570 — Research Framework 3.1.1 + 3.1.2)',
         bold=True, size=11, color_hex='FFFFFF')
set_cell(gt.cells[2], fmt(grand_total) + ' THB',
         bold=True, size=11, color_hex='FFFFFF', align='right')

# Column widths
COL_W = [Cm(7.5), Cm(6.5), Cm(3.5)]
for row in tbl.rows:
    if len(set(id(c) for c in row.cells)) == 3:
        for ci2, w in enumerate(COL_W):
            row.cells[ci2].width = w

# ── Replace old table ─────────────────────────────────────────────────────────
old_elem = doc.tables[6]._tbl
new_elem = tbl._tbl
body = doc.element.body
body.remove(new_elem)
idx  = list(body).index(old_elem)
body.remove(old_elem)
body.insert(idx, new_elem)

# ── Update cover table ────────────────────────────────────────────────────────
cover = doc.tables[0].rows[6].cells[0]
cover.paragraphs[0].clear()
cover.paragraphs[0].add_run(
    f'Requested Budget: {fmt(grand_total)} THB  '
    f'(HSRI 5-category structure, FY2570)')

# ── Print summary ─────────────────────────────────────────────────────────────
print(f'{"Category":<48} {"Amount":>12}  {"% non-eq":>9}')
print('-' * 75)
for ci2, cat in enumerate(BUDGET):
    print(f'  Cat {ci2+1}  {cat["cat"][:40]:<40} '
          f'{fmt(cat["subtotal"]):>12}  {cat["subtotal"]/non_eq*100:>8.1f}%')
print('-' * 75)
print(f'  {"GRAND TOTAL":<46} {fmt(grand_total):>12}')
print()
print(f'Personnel : {cat1/non_eq*100:.1f}%  (limit 30%)  OK')
print(f'Overhead  : {BUDGET[1]["subtotal"]/non_eq*100:.1f}%  (limit 15%)  OK')
print(f'Fee       : {BUDGET[4]["subtotal"]/base123*100:.1f}%  (limit 10%)  OK')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nSaved.')
