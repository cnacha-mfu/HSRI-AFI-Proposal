# -*- coding: utf-8 -*-
"""
Remove 'Policy dissemination workshop' row from Cat 2,
then cascade-update Cat 2 subtotal, Cat 5 item/subtotal,
compliance note, grand total, and cover table.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_EN_tmp.docx')

doc = Document(DOCX)
t   = doc.tables[6]

def fmt(n): return f'{int(n):,}'

def shade(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex); tcPr.append(shd)

def set_cell(cell, text, bold=False, italic=False, size=10,
             align='left', fill=None, color_hex=None):
    if fill: shade(cell, fill)
    para = cell.paragraphs[0]; para.clear()
    para.alignment = (WD_ALIGN_PARAGRAPH.RIGHT  if align == 'right'  else
                      WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else
                      WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16))

# ── 1. Remove the workshop row ────────────────────────────────────────────────
removed = False
for row in t.rows:
    if 'Policy dissemination workshop' in row.cells[0].text:
        t._tbl.remove(row._tr)
        print('Removed: Policy dissemination workshop row')
        removed = True
        break
if not removed:
    print('WARNING: row not found — check table index or text')

# ── 2. Updated totals ─────────────────────────────────────────────────────────
CAT1   = 961_200
CAT2   = 450_000          # was 500,000 — removed 50,000
CAT3   = 2_741_000
CAT4   = 400_000
base123 = CAT1 + CAT2 + CAT3   # 4,152,200
CAT5   = round(base123 * 0.10) # 415,220
GRAND  = CAT1 + CAT2 + CAT3 + CAT4 + CAT5
non_eq = GRAND - CAT4

pct1 = CAT1  / non_eq  * 100
pct2 = CAT2  / non_eq  * 100
pct5 = CAT5  / base123 * 100

# ── 3. Cascade updates (scan by text) ────────────────────────────────────────
for row in t.rows:
    txt0 = row.cells[0].text.strip()

    if 'Subtotal — Category 2' in txt0:
        set_cell(row.cells[0], '  Subtotal — Category 2', bold=True, size=10, fill='D5F5E3')
        set_cell(row.cells[1], '', size=10, fill='D5F5E3')
        set_cell(row.cells[2], fmt(CAT2), bold=True, size=10, align='right', fill='D5F5E3')
        print(f'Cat 2 subtotal → {fmt(CAT2)}')

    elif 'MFU institutional overhead' in txt0:
        set_cell(row.cells[1],
                 f'10% x (Category 1 + 2 + 3) = 10% x {fmt(base123)} THB',
                 italic=True, size=10, fill='EBF5EB', color_hex='444444')
        set_cell(row.cells[2], fmt(CAT5), size=10, align='right', fill='EBF5EB')
        print(f'Cat 5 item → {fmt(CAT5)}')

    elif 'Subtotal — Category 5' in txt0:
        set_cell(row.cells[0], '  Subtotal — Category 5', bold=True, size=10, fill='D5F5E3')
        set_cell(row.cells[1], '', size=10, fill='D5F5E3')
        set_cell(row.cells[2], fmt(CAT5), bold=True, size=10, align='right', fill='D5F5E3')
        print(f'Cat 5 subtotal → {fmt(CAT5)}')

    elif 'HSRI compliance' in txt0:
        msg = (f'HSRI compliance  |  '
               f'Personnel (Cat 1): {fmt(CAT1)} / {fmt(non_eq)} = {pct1:.1f}% <= 30%   '
               f'Overhead (Cat 2): {pct2:.1f}% <= 15%   '
               f'Fee (Cat 5): {pct5:.1f}% <= 10%')
        set_cell(row.cells[0], msg, italic=True, size=9, fill='FEF9E7')
        print('Compliance note updated')

    elif 'GRAND TOTAL' in txt0:
        set_cell(row.cells[2], fmt(GRAND) + ' THB',
                 bold=True, size=11, color_hex='FFFFFF', align='right', fill='0F6E56')
        print(f'Grand total → {fmt(GRAND)}')

# ── 4. Cover table ────────────────────────────────────────────────────────────
cover = doc.tables[0].rows[6].cells[0]
cover.paragraphs[0].clear()
cover.paragraphs[0].add_run(
    f'Requested Budget: {fmt(GRAND)} THB  '
    f'(HSRI 5-category structure, FY2570)')
print(f'Cover updated → {fmt(GRAND)} THB')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nSaved.')
print(f'\nCat1 {fmt(CAT1)} | Cat2 {fmt(CAT2)} | Cat3 {fmt(CAT3)} | Cat4 {fmt(CAT4)} | Cat5 {fmt(CAT5)}')
print(f'Grand total: {fmt(GRAND)} THB')
print(f'Personnel: {pct1:.1f}%  Overhead: {pct2:.1f}%  Fee: {pct5:.1f}%')
