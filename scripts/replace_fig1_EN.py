# -*- coding: utf-8 -*-
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

FOLDER  = r'G:\My Drive\Research\MORU'
DOCX    = os.path.join(FOLDER, 'Research_Proposal_EN.docx')
TMP     = os.path.join(FOLDER, 'Research_Proposal_EN_new.docx')
IMG_EN  = os.path.join(FOLDER, 'word', 'media', 'image1_EN.png')

doc = Document(DOCX)

# Find the paragraph holding the first embedded image
target_para = None
for i, para in enumerate(doc.paragraphs):
    if para._p.findall('.//' + qn('a:blip')):
        target_para = para
        print(f'Found image paragraph at index {i}: "{para.text[:40]}"')
        break

if target_para is None:
    print('ERROR: image paragraph not found'); sys.exit(1)

# Remove all runs (which contain the old image) from the paragraph
p_elem = target_para._p
for r in list(p_elem.findall(qn('w:r'))):
    p_elem.remove(r)

# Add the new English image as a run in the same paragraph
run = target_para.add_run()
run.add_picture(IMG_EN, width=Inches(6.2))
print('Inserted English Figure 1.')

doc.save(TMP)
print('Saved to:', TMP)

# Swap files
os.replace(TMP, DOCX)
print('Replaced original.')
