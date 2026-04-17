# -*- coding: utf-8 -*-
"""
Direct row-by-row Thai translation of the budget table in Research_Proposal_TH.docx.
Addresses all remaining English text by row index.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FOLDER = r'G:\My Drive\Research\MORU'
DOCX   = os.path.join(FOLDER, 'Research_Proposal_TH.docx')
TMP    = os.path.join(FOLDER, 'Research_Proposal_TH_tmp.docx')

doc = Document(DOCX)
t   = doc.tables[6]

def rewrite_para(para, text):
    """Replace text in a paragraph, preserving the first run's formatting."""
    if not para.runs:
        para.add_run(text)
        return
    f = para.runs[0]
    bold, italic, size, underline = f.bold, f.italic, f.font.size, f.underline
    color = None
    try:
        if f.font.color and f.font.color.type:
            color = f.font.color.rgb
    except Exception:
        pass
    para.clear()
    run = para.add_run(text)
    run.bold = bold; run.italic = italic; run.underline = underline
    if size: run.font.size = size
    if color: run.font.color.rgb = color

def fix(ri, ci, pi, text):
    """Fix a specific row/cell/paragraph."""
    row = t.rows[ri]
    # For merged rows, cells[0] == cells[1] == cells[2], so use cells[ci] directly
    unique = list(dict.fromkeys(id(c._tc) for c in row.cells))
    cell_map = {}
    idx = 0
    for c in row.cells:
        cid = id(c._tc)
        if cid not in cell_map:
            cell_map[cid] = idx
            idx += 1
    # Get the ci-th unique cell
    seen = []
    for c in row.cells:
        cid = id(c._tc)
        if cid not in [id(x._tc) for x in seen]:
            seen.append(c)
    if ci < len(seen):
        para = seen[ci].paragraphs[pi]
        print(f'  R{ri:02d}C{ci}P{pi}: {repr(para.text[:50])} → {repr(text[:50])}')
        rewrite_para(para, text)

# ── Header row ────────────────────────────────────────────────────────────────
fix(0, 0, 0, 'หมวดหมู่งบประมาณ / รายการ')

# ── Category 1 header ─────────────────────────────────────────────────────────
fix(1, 0, 0, 'หมวดที่ 1 — ค่าตอบแทนบุคลากร (หมวดค่าตอบแทน)')
fix(1, 0, 1, '21.0% ของยอดรวมไม่รวมครุภัณฑ์ — เป็นไปตามเพดาน สวรส. ไม่เกิน 30%')

# ── Cat 1 items ───────────────────────────────────────────────────────────────
fix(2, 0, 0, '  หัวหน้าโครงการวิจัย (ปริญญาเอก, รศ.)')
fix(2, 1, 0, '45,000 บาท/เดือน × 15% FTE × 18 เดือน')
fix(3, 1, 0, '35,000 บาท/เดือน × 19% FTE × 18 เดือน')
fix(4, 1, 0, '22,000 บาท/เดือน × 100% FTE × 18 เดือน')
fix(5, 0, 0, '  ผู้ช่วยวิจัยทางคลินิก (ปริญญาตรี)')
fix(5, 1, 0, '18,000 บาท/เดือน × 100% FTE × 18 เดือน')
fix(6, 0, 0, '  รวมย่อย — หมวดที่ 1')

# ── Category 2 header ─────────────────────────────────────────────────────────
fix(7, 0, 0, 'หมวดที่ 2 — ค่าบริหารจัดการโครงการ (หมวดค่าบริหารจัดการ)')
fix(7, 0, 1, '9.9% ของยอดรวมไม่รวมครุภัณฑ์ — เป็นไปตามเพดาน สวรส. ไม่เกิน 15%')

# ── Cat 2 items ───────────────────────────────────────────────────────────────
fix(8,  1, 0, '3 รายงาน × เหมาจ่าย')
fix(15, 0, 0, '  รวมย่อย — หมวดที่ 2')

# ── Category 3 header ─────────────────────────────────────────────────────────
fix(16, 0, 0, 'หมวดที่ 3 — ค่าดำเนินงานวิจัย (หมวดค่าดำเนินงาน)')
fix(16, 0, 1, 'รวมสัญญาช่วง MFU-MORU (ตามระเบียบ สวรส. ข้อ 3.2.3)')

# ── Cat 3 items ───────────────────────────────────────────────────────────────
fix(17, 1, 0, 'เฉลี่ย ~52,778 บาท/เดือน × 18 เดือน (รวมเบ็ดเสร็จ)')
fix(18, 0, 0, '  สัญญาช่วง CCRU/MORU — เงินเดือนผู้ช่วยวิจัยภาคสนาม [สัญญาช่วง MFU-MORU]')
fix(18, 1, 0, '25,000 บาท/เดือน × 12 เดือน')
fix(20, 0, 0, '  ค่าเดินทางเก็บข้อมูลภาคสนาม — ทีมวิจัย MFU')
fix(20, 1, 0, '20 ครั้งภาคสนาม × 10,000 บาท เฉลี่ย (เดินทาง + ที่พัก)')
fix(21, 0, 0, '  การอบรมเชิงปฏิบัติการสำหรับ อสม. และพยาบาล')
fix(21, 1, 0, '5 พื้นที่ฝึกอบรม × 24,000 บาท/พื้นที่ (เหมาจ่าย)')
fix(22, 0, 0, '  ค่าดำเนินการมีส่วนร่วมผู้เข้าร่วมและกระบวนการรับทราบยินยอม')
fix(22, 1, 0, '5 พื้นที่ × 16,000 บาท/พื้นที่ (ดำเนินการยินยอม + สิ่งจูงใจผู้เข้าร่วม เหมาจ่าย)')
fix(23, 0, 0, '  บริการถอดความและติดป้ายกำกับข้อมูล')
fix(23, 1, 0, 'ถอดเสียงทางคลินิก + ติดป้ายกำกับเอนทิตีทางการแพทย์; 2,000 บันทึก')
fix(25, 0, 0, '  ที่ปรึกษาผู้เชี่ยวชาญทางคลินิก — ผู้ทบทวนภายนอก 3 คน + ดำเนิน UAT')
fix(25, 1, 0, 'ตรวจสอบกฎ CDSS, ทบทวนโปรโตคอล UAT, ข้อเสนอแนะทางคลินิก 2 รอบ')
fix(26, 0, 0, '  ค่าสมัครสมาชิกซอฟต์แวร์ (LLM API, แพลตฟอร์ม Annotation, NLP, CI/CD)')
fix(26, 1, 0, 'สมัครสมาชิกหลายรายการ < 20,000 บาท/หน่วย; จัดเป็นค่าดำเนินงานตาม สวรส.')
fix(28, 0, 0, '  รวมย่อย — หมวดที่ 3')

# ── Category 4 header ─────────────────────────────────────────────────────────
fix(29, 0, 0, 'หมวดที่ 4 — ค่าครุภัณฑ์ (หมวดค่าครุภัณฑ์)')
fix(29, 0, 1, 'ทุกรายการ >= 20,000 บาท/หน่วย ไม่นำมาคำนวณสัดส่วนบุคลากร')

# ── Cat 4 items ───────────────────────────────────────────────────────────────
fix(32, 0, 0, '  เครื่องวัดสัญญาณชีพดิจิทัล (SpO2, ความดันโลหิต, อุณหภูมิ) — ทดสอบภาคสนาม')
fix(33, 0, 0, '  รวมย่อย — หมวดที่ 4')

# ── Category 5 header ─────────────────────────────────────────────────────────
fix(34, 0, 0, 'หมวดที่ 5 — ค่าธรรมเนียมสถาบัน (หมวดค่าธรรมเนียมหน่วยงานผู้รับทุน)')
fix(34, 0, 1, '10% ของ (หมวด 1+2+3) = 4,152,200 บาท — เป็นไปตามเพดาน สวรส. ไม่เกิน 10%')

# ── Cat 5 items ───────────────────────────────────────────────────────────────
fix(35, 0, 0, '  ค่าธรรมเนียมสถาบัน MFU — สาธารณูปโภค พื้นที่ทำงาน ธุรการ โครงสร้างพื้นฐาน IT')
fix(35, 1, 0, '10% × (หมวด 1 + 2 + 3) = 10% × 4,152,200 บาท')
fix(36, 0, 0, '  รวมย่อย — หมวดที่ 5')

# ── Compliance note ───────────────────────────────────────────────────────────
fix(37, 0, 0,
    'การตรวจสอบการปฏิบัติตาม สวรส.  |  '
    'บุคลากร (หมวด 1): 961,200 / 4,567,420 = 21.0% ≤ 30%   '
    'ค่าบริหาร (หมวด 2): 9.9% ≤ 15%   '
    'ค่าธรรมเนียม (หมวด 5): 10.0% ≤ 10%')

doc.save(TMP)
os.replace(TMP, DOCX)
print('\nDone — saved.')
